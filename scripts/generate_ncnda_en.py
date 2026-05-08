"""
scripts/generate_ncnda_en.py — v3
==================================
Gera o NCNDA EN com design idêntico ao HTML/PDF modelo Samba:
  - Capa escura (fundo #1A1A1A, logo, título, metadata)
  - Tabela de Partes (3 colunas, cabeçalhos escuros, campos com linhas)
  - "CLAUSE N" em laranja + título em negrito (sem tabela extra)
  - Caixas highlight (tabela com borda esquerda laranja)
  - Bloco de assinaturas (tabela 3 colunas)
  - Footer bar escuro

Faz upload para Google Drive convertendo para Google Docs.
"""
from __future__ import annotations

import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from dotenv import load_dotenv
load_dotenv(override=True)

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta de cores ───────────────────────────────────────────────────────────
ORANGE  = "FA8200"
BLACK   = "1A1A1A"
WHITE   = "FFFFFF"
GRAY    = "888888"
LGRAY   = "DDDDDD"
TEXT    = "2D2D2D"
TEXTLT  = "555555"
BGLIGHT = "FFF8F0"   # highlight box background

# ── Tipografia ────────────────────────────────────────────────────────────────
FONT    = "Montserrat"
FONT_SM = Pt(9)
FONT_MD = Pt(10.5)
FONT_LG = Pt(13)
FONT_XL = Pt(20)

FOLDER_ID         = os.getenv("GOOGLE_DRIVE_FOLDER_ID", "")
NCNDA_FOLDER_NAME = "__NCDA"


# ── Helpers XML / formatação ──────────────────────────────────────────────────

def _hex_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _cell_bg(cell, hex_color: str):
    """Define a cor de fundo de uma célula."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove shd anterior
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper().lstrip("#"))
    tcPr.append(shd)


def _cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"),    str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _no_borders(table):
    """Remove todas as bordas de uma tabela."""
    tbl   = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _left_border(cell, color: str = ORANGE, sz: int = 18):
    """Define apenas a borda esquerda de uma célula (para highlight box)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for side in ("top", "right", "bottom", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"),  "0")
        tcB.append(b)
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(sz))
    left.set(qn("w:color"), color.lstrip("#").upper())
    tcB.append(left)
    tcPr.append(tcB)


def _table_width_pct(table, pct: int = 100):
    tbl   = table._tbl
    tblPr = tbl.tblPr
    tblW  = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(pct * 50))
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


def _run(para, text: str, bold: bool = False, italic: bool = False,
         color: str | None = None, size: Pt | None = None,
         font: str | None = None) -> object:
    run = para.add_run(text)
    run.font.name  = font or FONT
    run.font.size  = size or FONT_MD
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = _hex_rgb(color)
    else:
        run.font.color.rgb = _hex_rgb(TEXT)
    return run


def _para(doc: Document, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_before: Pt = Pt(0), space_after: Pt = Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after  = space_after
    return p


def _blank(doc: Document, n: int = 1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)


def _cell_para(cell, align=WD_ALIGN_PARAGRAPH.LEFT) -> object:
    """Retorna o parágrafo padrão de uma célula (ou adiciona um)."""
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def _cell_run(cell, text: str, bold: bool = False, color: str = WHITE,
              size: Pt = FONT_MD, align=WD_ALIGN_PARAGRAPH.LEFT, font: str = FONT) -> object:
    p   = _cell_para(cell, align)
    run = p.add_run(text)
    run.font.name      = font
    run.font.size      = size
    run.bold           = bold
    run.font.color.rgb = _hex_rgb(color)
    return run


def _cell_add_line(cell, text: str, bold: bool = False, color: str = TEXT,
                   size: Pt = FONT_MD, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Adiciona uma linha extra de texto em uma célula."""
    p   = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.font.name      = FONT
    run.font.size      = size
    run.bold           = bold
    run.font.color.rgb = _hex_rgb(color)
    return run


# ── Cláusula helper ───────────────────────────────────────────────────────────

def _clause_label(doc: Document, number: str):
    """'CLAUSE N' em laranja, pequeno, bold."""
    p = _para(doc, WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(12))
    _run(p, f"CLAUSE {number}", bold=True, color=ORANGE, size=FONT_SM)


def _clause_title(doc: Document, title: str):
    """Título da cláusula em negrito grande."""
    p = _para(doc, WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6))
    _run(p, title, bold=True, size=Pt(13), color=BLACK)


def _body(doc: Document, parts: list[tuple[str, bool]],
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm: float = 0.0):
    """Parágrafo de corpo com múltiplos runs [(texto, bold), ...]."""
    p = _para(doc, align)
    if indent_cm:
        from docx.shared import Cm as Cm_
        p.paragraph_format.left_indent = Cm_(indent_cm)
    for text, bold in parts:
        _run(p, text, bold=bold, size=FONT_MD)
    return p


def _list_line(doc: Document, text: str):
    """Linha de lista com recuo."""
    p = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    _run(p, text, size=FONT_MD)


def _highlight_box(doc: Document, label: str, content: str):
    """Caixa highlight com borda esquerda laranja (#FFF8F0 background)."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    _table_width_pct(tbl, 100)
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, BGLIGHT)
    _left_border(cell, ORANGE, sz=24)
    _cell_margins(cell, top=100, bottom=100, left=150, right=150)

    p1 = _cell_para(cell)
    r1 = p1.add_run(label)
    r1.font.name = FONT; r1.font.size = FONT_MD; r1.bold = True
    r1.font.color.rgb = _hex_rgb(BLACK)
    r2 = p1.add_run(content)
    r2.font.name = FONT; r2.font.size = FONT_MD; r2.bold = False
    r2.font.color.rgb = _hex_rgb(TEXT)


def _divider(doc: Document):
    """Linha horizontal (tabela de 1 px de altura com borda inferior)."""
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tcB.append(b)
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:color"), LGRAY)
    tcB.append(bot)
    tcPr.append(tcB)
    row = tbl.rows[0]
    row.height = Pt(2)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    _table_width_pct(tbl, 100)


# ── Seções do documento ───────────────────────────────────────────────────────

def _build_cover(doc: Document):
    """Capa escura: fundo #1A1A1A, logo, título, metadata grid."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    _table_width_pct(tbl, 100)
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, BLACK)
    _cell_margins(cell, top=700, bottom=700, left=700, right=700)

    # Row height grande para simular página inteira
    row = tbl.rows[0]
    row.height = Cm(20)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # Logo: "sambaEXPORT"
    p_logo = _cell_para(cell, WD_ALIGN_PARAGRAPH.LEFT)
    r_s = p_logo.add_run("samba")
    r_s.font.name = FONT; r_s.font.size = Pt(22); r_s.bold = True
    r_s.font.color.rgb = _hex_rgb(ORANGE)
    r_e = p_logo.add_run("EXPORT")
    r_e.font.name = FONT; r_e.font.size = Pt(22); r_e.bold = True
    r_e.font.color.rgb = _hex_rgb(WHITE)

    # Espaço
    for _ in range(3):
        pb = cell.add_paragraph()
        pb.paragraph_format.space_before = Pt(0)
        pb.paragraph_format.space_after  = Pt(0)

    # Título
    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(2)
    rt = p_title.add_run("Non-Circumvention,\nNon-Disclosure &\n")
    rt.font.name = FONT; rt.font.size = Pt(22); rt.bold = True
    rt.font.color.rgb = _hex_rgb(WHITE)
    rt2 = p_title.add_run("Confidentiality Agreement")
    rt2.font.name = FONT; rt2.font.size = Pt(22); rt2.bold = True
    rt2.font.color.rgb = _hex_rgb(ORANGE)

    # Subtítulo
    p_sub = cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_before = Pt(4)
    p_sub.paragraph_format.space_after  = Pt(20)
    rs = p_sub.add_run("NCNDA — International Commercial Intermediation")
    rs.font.name = FONT; rs.font.size = FONT_SM
    rs.font.color.rgb = _hex_rgb(GRAY)

    # Linha laranja decorativa
    p_line = cell.add_paragraph()
    r_line = p_line.add_run("━━━━━━━━━━━━━━━")
    r_line.font.name = FONT; r_line.font.size = Pt(8)
    r_line.font.color.rgb = _hex_rgb(ORANGE)

    # Metadados 2x3 grid (simulado com tabs)
    meta = [
        ("DATE OF EXECUTION",              "DD / MM / 2026"),
        ("DOCUMENT REFERENCE",             "NCNDA-SAMBA-2026-___"),
        ("PARTY I — ORIGINATOR",           "SAMBA EXPORT LTDA"),
        ("GOVERNING LAW",                  "Brazilian Law — São Paulo Courts"),
        ("PARTY II — BRAZILIAN INTERMEDIARY", "Company name to be filled"),
        ("PARTY III — BUYER-SIDE INTERMEDIARY", "Company name to be filled"),
    ]
    for label, val in meta:
        pb = cell.add_paragraph()
        pb.paragraph_format.space_before = Pt(8)
        pb.paragraph_format.space_after  = Pt(0)
        rl = pb.add_run(label + "\n")
        rl.font.name = FONT; rl.font.size = Pt(7.5); rl.bold = True
        rl.font.color.rgb = _hex_rgb(GRAY)
        rv = pb.add_run(val)
        rv.font.name = FONT; rv.font.size = Pt(11); rv.bold = True
        rv.font.color.rgb = _hex_rgb(WHITE)

    # CONFIDENTIAL watermark
    p_conf = cell.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_conf.paragraph_format.space_before = Pt(20)
    rc = p_conf.add_run("CONFIDENTIAL")
    rc.font.name = FONT; rc.font.size = FONT_SM; rc.bold = True
    rc.font.color.rgb = _hex_rgb("333333")

    # Page break após a capa
    doc.add_page_break()


def _build_doc_header(doc: Document):
    """Cabeçalho do documento: logo esquerda | info direita, borda laranja inferior."""
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    _table_width_pct(tbl, 100)

    # Borda inferior laranja
    for cell in tbl.rows[0].cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcB  = OxmlElement("w:tcBorders")
        for side in ("top", "left", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tcB.append(b)
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "12")
        bot.set(qn("w:color"), ORANGE)
        tcB.append(bot)
        tcPr.append(tcB)

    # Logo esquerda
    cell_l = tbl.rows[0].cells[0]
    _cell_margins(cell_l, top=120, bottom=120, left=0, right=0)
    p_l = _cell_para(cell_l)
    r1 = p_l.add_run("samba")
    r1.font.name = FONT; r1.font.size = Pt(16); r1.bold = True
    r1.font.color.rgb = _hex_rgb(ORANGE)
    r2 = p_l.add_run("EXPORT")
    r2.font.name = FONT; r2.font.size = Pt(16); r2.bold = True
    r2.font.color.rgb = _hex_rgb(BLACK)

    # Info direita
    cell_r = tbl.rows[0].cells[1]
    _cell_margins(cell_r, top=100, bottom=100, left=0, right=0)
    p_r = _cell_para(cell_r, WD_ALIGN_PARAGRAPH.RIGHT)
    ri = p_r.add_run(
        "NCNDA — Non-Circumvention, Non-Disclosure & Confidentiality Agreement\n"
        "Ref: NCNDA-SAMBA-2026-___  |  Date: ____/____/2026"
    )
    ri.font.name = FONT; ri.font.size = Pt(7.5)
    ri.font.color.rgb = _hex_rgb(TEXTLT)

    _blank(doc)


def _build_title_block(doc: Document):
    """Bloco de título centrado."""
    p = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    r1 = _run(p, "Non-Circumvention, Non-Disclosure &\n", bold=True, size=Pt(18), color=BLACK)
    r2 = _run(p, "Confidentiality Agreement", bold=True, size=Pt(18), color=ORANGE)

    p2 = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))
    _run(p2, "NCNDA — International Agricultural Commodity Intermediation",
         size=FONT_SM, color=TEXTLT)


def _build_preamble(doc: Document):
    """Parágrafo de abertura."""
    p = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(10))
    _run(p, 'This Non-Circumvention, Non-Disclosure and Confidentiality Agreement (')
    _run(p, '"Agreement"', bold=True)
    _run(p, ') is entered into as of ')
    _run(p, '______ / ______ / 2026', bold=True)
    _run(p, ' by and among the following parties (individually a ')
    _run(p, '"Party"', bold=True)
    _run(p, ' and collectively the ')
    _run(p, '"Parties"', bold=True)
    _run(p, '):')


def _build_parties_table(doc: Document):
    """Tabela de partes: 3 colunas com cabeçalhos escuros e campos."""

    parties = [
        {
            "header": "PARTY I — ORIGINATOR",
            "fields": [
                ("LEGAL NAME",            "SAMBA EXPORT LTDA",                    True),
                ("BRAZILIAN TAX ID (CNPJ)","60.280.015/0001-82",                  True),
                ("REGISTERED ADDRESS",    "Av. Brigadeiro Faria Lima, 1811, Suite 115\nSão Paulo, SP — Brazil", True),
                ("LEGAL REPRESENTATIVE",  "Marcelo Soares Magalhães Nogueira",    True),
                ("ROLE IN AGREEMENT",     "Brazilian Commodity Originator & Exporter", True),
            ],
        },
        {
            "header": "PARTY II — BR INTERMEDIARY",
            "fields": [
                ("LEGAL NAME",            "____________________________", False),
                ("BRAZILIAN TAX ID (CNPJ)","____________________________", False),
                ("REGISTERED ADDRESS",    "____________________________", False),
                ("LEGAL REPRESENTATIVE",  "____________________________", False),
                ("PASSPORT / ID NUMBER",  "____________________________", False),
            ],
        },
        {
            "header": "PARTY III — BUYER SIDE",
            "fields": [
                ("LEGAL NAME",            "____________________________", False),
                ("COUNTRY OF INCORPORATION", "____________________________", False),
                ("COMPANY REGISTRATION NO.", "____________________________", False),
                ("REGISTERED ADDRESS",    "____________________________", False),
                ("LEGAL REPRESENTATIVE",  "____________________________", False),
                ("PASSPORT NUMBER",       "____________________________", False),
            ],
        },
    ]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_width_pct(tbl, 100)

    # Bordas finas cinza
    tblB = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:color"), LGRAY)
        tblB.append(b)
    tbl._tbl.tblPr.append(tblB)

    # Linha de cabeçalhos
    hrow = tbl.rows[0]
    for i, party in enumerate(parties):
        cell = hrow.cells[i]
        _cell_bg(cell, BLACK)
        _cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p_h = _cell_para(cell)

        # dot laranja + nome
        r_dot = p_h.add_run("● ")
        r_dot.font.name = FONT; r_dot.font.size = FONT_SM; r_dot.bold = True
        r_dot.font.color.rgb = _hex_rgb(ORANGE)
        r_lbl = p_h.add_run(party["header"])
        r_lbl.font.name = FONT; r_lbl.font.size = FONT_SM; r_lbl.bold = True
        r_lbl.font.color.rgb = _hex_rgb(WHITE)

    # Determina nº máximo de campos
    max_fields = max(len(p["fields"]) for p in parties)

    # Adiciona linhas de campos
    for fi in range(max_fields):
        row = tbl.add_row()
        for ci, party in enumerate(parties):
            cell = row.cells[ci]
            _cell_margins(cell, top=100, bottom=80, left=120, right=120)

            if fi < len(party["fields"]):
                label, value, is_fixed = party["fields"][fi]

                # Label pequeno cinza
                p_lbl = _cell_para(cell)
                r_l = p_lbl.add_run(label)
                r_l.font.name = FONT; r_l.font.size = Pt(7.5); r_l.bold = True
                r_l.font.color.rgb = _hex_rgb(GRAY)

                # Valor
                p_val = cell.add_paragraph()
                p_val.paragraph_format.space_before = Pt(2)
                p_val.paragraph_format.space_after  = Pt(0)
                r_v = p_val.add_run(value)
                r_v.font.name = FONT
                r_v.font.size = FONT_SM
                r_v.bold      = is_fixed
                r_v.font.color.rgb = _hex_rgb(BLACK if is_fixed else "BBBBBB")

    _blank(doc)


# ── Corpo do documento ────────────────────────────────────────────────────────

def build_document() -> Document:
    doc = Document()

    # Margens A4
    section = doc.sections[0]
    section.page_width   = Cm(21.0)
    section.page_height  = Cm(29.7)
    section.top_margin   = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── CAPA ───────────────────────────────────────────────────────────────────
    _build_cover(doc)

    # ── CABEÇALHO ─────────────────────────────────────────────────────────────
    _build_doc_header(doc)
    _blank(doc)

    # ── TÍTULO ────────────────────────────────────────────────────────────────
    _build_title_block(doc)

    # ── PREAMBLE ──────────────────────────────────────────────────────────────
    _build_preamble(doc)

    # ── TABELA DE PARTES ──────────────────────────────────────────────────────
    _build_parties_table(doc)

    _divider(doc)
    _blank(doc)

    # ── CLAUSE 1 ──────────────────────────────────────────────────────────────
    _clause_label(doc, "1")
    _clause_title(doc, "Purpose")

    _body(doc, [
        ("The purpose of this Agreement is to formalize the collaborative intermediation "
         "relationship among the Parties in connection with the purchase, sale, and export of "
         "agricultural commodities of Brazilian origin — including but not limited to soybeans, "
         "soybean meal, soybean oil, corn, sugar (ICUMSA 45, VHP, Crystal), coffee, cotton, and "
         "animal proteins — as well as any other business arising directly or indirectly from "
         "introductions made under this Agreement.", False),
    ])

    _blank(doc)

    _body(doc, [
        ('The Parties acknowledge that each brings to this collaboration confidential information, '
         'business relationships, and strategic contacts (', False),
        ('"Protected Contacts"', True),
        (') that have intrinsic commercial value and must be protected accordingly.', False),
    ])

    # ── CLAUSE 2 ──────────────────────────────────────────────────────────────
    _clause_label(doc, "2")
    _clause_title(doc, "Role of the Parties")

    _body(doc, [
        ("Each Party acts exclusively as a ", False),
        ("commercial intermediary", True),
        (" in the transactions facilitated under this Agreement. No Party shall be deemed "
         "a principal buyer, seller, or guarantor of any commodity transaction.", False),
    ])

    _blank(doc)
    _body(doc, [("Accordingly, no Party assumes any liability for:", False)])

    for item in [
        "Product quality, specifications, or conformity;",
        "Logistics, shipping, or customs clearance;",
        "Financial performance or payment by end buyers;",
        "Force majeure events affecting the underlying commodity transaction.",
    ]:
        _list_line(doc, f"• {item}")

    _blank(doc)
    _body(doc, [
        ("All liability for the underlying commercial transaction rests solely with the "
         "parties to the main Sales & Purchase Agreement (SPA).", False),
    ])

    # ── CLAUSE 3 ──────────────────────────────────────────────────────────────
    _clause_label(doc, "3")
    _clause_title(doc, "Commission Entitlement")

    _body(doc, [
        ("Each Party shall be entitled to receive a commission for its role in facilitating "
         "transactions under this Agreement. Commission amounts, rates, and split arrangements "
         "shall be agreed on a per-deal basis and may be formalized via:", False),
    ])

    for item in [
        "Written addendum to this Agreement; or",
        "Confirming message exchanged via WhatsApp, Telegram, WeChat, or corporate email — "
        "which shall have the same legal force as a written addendum.",
    ]:
        _list_line(doc, f"• {item}")

    _blank(doc)
    _highlight_box(doc,
        "Payment timing: ",
        "All commissions shall be paid within 1 (one) business day of confirmed receipt of "
        "cleared funds in the designated bank account of the paying Party.")
    _blank(doc)

    _body(doc, [
        ("Each Party is solely responsible for the payment of all applicable taxes, duties, "
         "and levies on its respective commission income under the laws of its jurisdiction.", False),
    ])

    # ── CLAUSE 4 ──────────────────────────────────────────────────────────────
    _clause_label(doc, "4")
    _clause_title(doc, "Commission Distribution")

    _body(doc, [
        ("The total gross commission generated from any transaction intermediated under this "
         "Agreement shall be distributed among the Parties according to the percentages or "
         "amounts agreed for each specific deal, as formalized under Clause 3 above.", False),
    ])

    _blank(doc)
    _body(doc, [
        ("No Party shall be entitled to alter the agreed distribution without the written "
         "consent of all Parties.", False),
    ])

    # ── CLAUSE 5 ──────────────────────────────────────────────────────────────
    _clause_label(doc, "5")
    _clause_title(doc, "Confidentiality, Non-Circumvention & Non-Compete")

    sub5 = [
        ("5.1", "Confidentiality.",
         "Each Party agrees to keep strictly confidential all information received from the "
         "other Parties, including but not limited to: business contacts, pricing, terms of "
         "deals, buyer and seller identities, and internal strategy. Confidential information "
         "shall not be disclosed to any third party without prior written consent."),

        ("5.2", "Non-Circumvention.",
         "Each Party irrevocably undertakes not to directly or indirectly contact, negotiate "
         "with, or transact with any Protected Contact introduced by another Party, without "
         "that Party's prior written authorization and without ensuring that Party's commission "
         "is properly secured."),
    ]

    for num, label, text in sub5:
        p = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
        p.paragraph_format.left_indent = Cm(0.8)
        _run(p, f"{num}  ", bold=True, color=ORANGE)
        _run(p, label + " ", bold=True)
        _run(p, text)
        _blank(doc)

    # 5.3 com sub-items
    p53 = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p53.paragraph_format.left_indent = Cm(0.8)
    _run(p53, "5.3  ", bold=True, color=ORANGE)
    _run(p53, "Consequences of Circumvention. ", bold=True)
    _run(p53, "Any Party found to have circumvented another Party shall be liable for:")

    sub53 = [
        ("5.3.1", "Full payment of the commission that would have been due to the circumvented "
                  "Party on the bypassed transaction;"),
        ("5.3.2", "All direct and indirect damages proven to have resulted from the circumvention;"),
        ("5.3.3", "A contractual penalty of _______________________________________________"
                  " (to be agreed by the Parties and inserted here), without prejudice to "
                  "additional damages."),
    ]

    for num2, text2 in sub53:
        p2 = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
        p2.paragraph_format.left_indent = Cm(1.5)
        _run(p2, f"{num2}  ", bold=True, color=ORANGE)
        _run(p2, text2)

    _blank(doc)

    # 5.4 Non-Compete
    p54 = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p54.paragraph_format.left_indent = Cm(0.8)
    _run(p54, "5.4  ", bold=True, color=ORANGE)
    _run(p54, "Non-Compete. ", bold=True)
    _run(p54, "No Party shall, for a period of ")
    _run(p54, "2 (two) years", bold=True)
    _run(p54, " after the termination of this Agreement, directly or indirectly contact, "
         "solicit, or transact with any Protected Contact introduced under this Agreement, "
         "without the introducing Party's prior written consent.")

    _blank(doc)

    # 5.5 Survival
    p55 = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p55.paragraph_format.left_indent = Cm(0.8)
    _run(p55, "5.5  ", bold=True, color=ORANGE)
    _run(p55, "Survival. ", bold=True)
    _run(p55, "The obligations under this Clause 5 shall survive termination of this "
         "Agreement for a period of ")
    _run(p55, "5 (five) years", bold=True)
    _run(p55, " (confidentiality and non-circumvention) and ")
    _run(p55, "2 (two) years", bold=True)
    _run(p55, " (non-compete).")

    _blank(doc)

    # ── CLAUSE 6 ──────────────────────────────────────────────────────────────
    _clause_label(doc, "6")
    _clause_title(doc, "Governing Law & Jurisdiction")

    _body(doc, [
        ("This Agreement shall be governed by and construed in accordance with the laws of the ", False),
        ("Federative Republic of Brazil", True),
        (".", False),
    ])

    _blank(doc)

    _body(doc, [
        ("Any dispute arising out of or in connection with this Agreement — including disputes "
         "regarding its formation, validity, breach, termination, or interpretation — shall be "
         "submitted to the exclusive jurisdiction of the ", False),
        ("Courts of the City of São Paulo, State of São Paulo, Brazil", True),
        (", with each Party irrevocably waiving any objection to such jurisdiction.", False),
    ])

    _blank(doc)

    _highlight_box(doc, "",
        "The Parties acknowledge that this Agreement may be entered into and enforced across "
        "multiple jurisdictions and that the choice of Brazilian law is deliberate and binding.")

    _blank(doc)

    # ── CLAUSE 7 ──────────────────────────────────────────────────────────────
    _clause_label(doc, "7")
    _clause_title(doc, "Term")

    _body(doc, [
        ("This Agreement shall remain in full force and effect for as long as business "
         "operations arising from introductions made hereunder are ongoing. The commission "
         "and confidentiality obligations shall survive any termination or expiry of this "
         "Agreement as set forth in Clause 5.5.", False),
    ])

    # ── CLAUSE 8 ──────────────────────────────────────────────────────────────
    _clause_label(doc, "8")
    _clause_title(doc, "Electronic Communications as Binding Agreements")

    _body(doc, [
        ("The Parties expressly agree that messages exchanged via ", False),
        ("WhatsApp, Telegram, WeChat, or corporate email", True),
        (" that confirm commission values, deal terms, or specific amendments to this "
         "Agreement shall constitute legally binding addenda to this Agreement, enforceable "
         "as if made in writing.", False),
    ])

    _blank(doc)

    _body(doc, [
        ("No Party shall contest the enforceability of such communications on the grounds "
         "of form or formality.", False),
    ])

    # ── CLAUSE 9 ──────────────────────────────────────────────────────────────
    _clause_label(doc, "9")
    _clause_title(doc, "Electronic Signature")

    _body(doc, [
        ("This Agreement may be executed electronically using any of the following platforms: ", False),
        ("DocuSign, ClickSign, Adobe Sign, Gov.br, or ICP-Brasil", True),
        (". Electronically signed versions shall constitute the original and shall have full "
         "legal force and effect under Brazilian law.", False),
    ])

    _blank(doc)

    _body(doc, [
        ("Each Party shall receive a signed copy upon completion of the signature process.", False),
    ])

    _blank(doc)
    _divider(doc)
    _blank(doc)

    # ── ASSINATURAS ───────────────────────────────────────────────────────────
    p_sig_title = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    _run(p_sig_title,
         "SIGNATURES — IN WITNESS WHEREOF, THE PARTIES HAVE EXECUTED THIS AGREEMENT",
         bold=True, size=FONT_SM, color=TEXTLT)

    sig_tbl = doc.add_table(rows=1, cols=3)
    _table_width_pct(sig_tbl, 100)

    # Bordas: só borda superior escura em cada bloco
    for ci, (party_label, party_name, rep) in enumerate([
        ("PARTY I — ORIGINATOR", "SAMBA EXPORT LTDA",
         "Marcelo Soares Magalhães Nogueira\nLegal Representative"),
        ("PARTY II — BR INTERMEDIARY", "____________________________",
         "____________________________\nLegal Representative"),
        ("PARTY III — BUYER SIDE", "____________________________",
         "____________________________\nLegal Representative"),
    ]):
        cell = sig_tbl.rows[0].cells[ci]
        _cell_margins(cell, top=100, bottom=100, left=80, right=80)

        # Borda superior
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcB  = OxmlElement("w:tcBorders")
        for side in ("bottom", "left", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tcB.append(b)
        top_b = OxmlElement("w:top")
        top_b.set(qn("w:val"),   "single")
        top_b.set(qn("w:sz"),    "12")
        top_b.set(qn("w:color"), BLACK)
        tcB.append(top_b)
        tcPr.append(tcB)

        # Label laranja
        p0 = _cell_para(cell)
        r0 = p0.add_run(party_label)
        r0.font.name = FONT; r0.font.size = Pt(7.5); r0.bold = True
        r0.font.color.rgb = _hex_rgb(ORANGE)

        # Nome da parte
        p1 = cell.add_paragraph()
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after  = Pt(2)
        r1 = p1.add_run(party_name)
        r1.font.name = FONT; r1.font.size = FONT_SM; r1.bold = True
        r1.font.color.rgb = _hex_rgb(BLACK)

        # Representante
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(14)
        r2 = p2.add_run(rep)
        r2.font.name = FONT; r2.font.size = FONT_SM
        r2.font.color.rgb = _hex_rgb(TEXTLT)

        # Linha de assinatura
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_before = Pt(28)
        p3.paragraph_format.space_after  = Pt(2)
        r3 = p3.add_run("_" * 30)
        r3.font.name = FONT; r3.font.size = FONT_SM
        r3.font.color.rgb = _hex_rgb("CCCCCC")

        # "Signature & Date"
        p4 = cell.add_paragraph()
        p4.paragraph_format.space_before = Pt(2)
        r4 = p4.add_run("Signature & Date")
        r4.font.name = FONT; r4.font.size = Pt(7.5)
        r4.font.color.rgb = _hex_rgb(GRAY)

    _blank(doc)
    _blank(doc)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    ftbl = doc.add_table(rows=1, cols=3)
    _table_width_pct(ftbl, 100)
    _no_borders(ftbl)
    frow = ftbl.rows[0]
    for cell in frow.cells:
        _cell_bg(cell, BLACK)
        _cell_margins(cell, top=100, bottom=100, left=120, right=120)

    # Logo
    fl = frow.cells[0]
    pf1 = _cell_para(fl)
    rf1 = pf1.add_run("samba")
    rf1.font.name = FONT; rf1.font.size = Pt(9); rf1.bold = True
    rf1.font.color.rgb = _hex_rgb(ORANGE)
    rf2 = pf1.add_run("EXPORT — CONFIDENTIAL")
    rf2.font.name = FONT; rf2.font.size = Pt(9); rf2.bold = True
    rf2.font.color.rgb = _hex_rgb("888888")

    # Ref
    fc = frow.cells[1]
    pf2 = _cell_para(fc, WD_ALIGN_PARAGRAPH.CENTER)
    rf3 = pf2.add_run("NCNDA-SAMBA-2026-___  |  Brazilian Law — São Paulo")
    rf3.font.name = FONT; rf3.font.size = Pt(7.5)
    rf3.font.color.rgb = _hex_rgb("555555")

    # Site
    fr = frow.cells[2]
    pf3 = _cell_para(fr, WD_ALIGN_PARAGRAPH.RIGHT)
    rf4 = pf3.add_run("sambaexport.com.br")
    rf4.font.name = FONT; rf4.font.size = Pt(7.5)
    rf4.font.color.rgb = _hex_rgb("555555")

    return doc


# ── Upload para Google Drive como Google Docs ─────────────────────────────────

def find_ncda_folder(drive_service) -> str | None:
    """Encontra o folder ID da pasta __NCDA no Drive."""
    results = drive_service.files().list(
        q=f"name='{NCNDA_FOLDER_NAME}' and mimeType='application/vnd.google-apps.folder' and trashed=false",
        fields="files(id, name)",
        supportsAllDrives=True,
        includeItemsFromAllDrives=True,
    ).execute()
    files = results.get("files", [])
    if files:
        return files[0]["id"]
    return None


def delete_existing(drive_service, folder_id: str | None) -> None:
    """Remove versão anterior do documento se existir."""
    q = "name='NCNDA - SAMBA INTERM DE NEGOCIOS (EN)' and mimeType='application/vnd.google-apps.document' and trashed=false"
    if folder_id:
        q += f" and '{folder_id}' in parents"
    results = drive_service.files().list(
        q=q, fields="files(id, name)",
        supportsAllDrives=True,
        includeItemsFromAllDrives=True,
    ).execute()
    for f in results.get("files", []):
        try:
            drive_service.files().delete(fileId=f["id"], supportsAllDrives=True).execute()
            print(f"  [DEL] versão anterior removida: {f['id']}")
        except Exception as e:
            print(f"  [AVISO] Não foi possível remover versão anterior ({f['id']}): {e}")


def upload_to_drive(doc: Document) -> str | None:
    from services.google_drive import drive_manager
    from googleapiclient.http import MediaIoBaseUpload

    if not drive_manager.service:
        print("  [ERRO] Drive não autenticado.")
        return None

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    folder_id = find_ncda_folder(drive_manager.service)
    if not folder_id:
        print("  [AVISO] Pasta __NCDA não encontrada — usando pasta raiz SAMBA.")
        folder_id = FOLDER_ID or None

    # Remove versão anterior para evitar duplicatas
    delete_existing(drive_manager.service, folder_id)

    file_metadata = {
        "name": "NCNDA - SAMBA INTERM DE NEGOCIOS (EN)",
        "mimeType": "application/vnd.google-apps.document",
        "parents": [folder_id] if folder_id else [],
    }

    media = MediaIoBaseUpload(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=False,
    )

    created = drive_manager.service.files().create(
        body=file_metadata,
        media_body=media,
        fields="id, webViewLink, name",
        supportsAllDrives=True,
    ).execute()

    url  = created.get("webViewLink")
    name = created.get("name")
    print(f"  [OK] Google Doc criado: '{name}'")
    print(f"  URL: {url}")
    return url


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print("=== NCNDA EN — v3 (design Samba) ===\n")

    print("1. Construindo documento...")
    doc = build_document()
    print(f"   {len(doc.paragraphs)} parágrafos gerados")

    local_path = (
        r"H:\Drives compartilhados\SAMBA EXPORT\MODELOS DE DOCUMENTOS"
        r"\__NCDA\NCNDA - SAMBA INTERM DE NEGOCIOS (EN).docx"
    )
    doc.save(local_path)
    print(f"   Backup local salvo: {local_path}")

    print("\n2. Upload para Google Drive como Google Docs...")
    url = upload_to_drive(doc)

    if url:
        print(f"\n=== CONCLUÍDO ===")
        print(f"Google Doc: {url}")
    else:
        print("\n[AVISO] Upload falhou — use o .docx local como fallback.")


if __name__ == "__main__":
    main()
