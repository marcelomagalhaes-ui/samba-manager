# -*- coding: utf-8 -*-
"""
services/loi_template_engine.py
================================
Engine de renderização de LOIs — alinhado aos templates REAIS (Caminho A).

Padrões reconhecidos nos templates:
  1. Chaves universais       {CITY}, {YYYY}, {DESTINATARY_LOIFULLNAME}, {XXX,XXX}, etc.
  2. Marcadores de produto   {Yellow Soybean, GMO - Grade #2}, {1201.90 (Yellow Soybean)},
                              {Yellow Soybean Parameters: ...}
                              → matched por strict/family keywords do produto
                              → o do produto ESCOLHIDO é mantido (sem braces)
                              → os de OUTROS produtos são removidos
                                  (linha de tabela inteira é removida se
                                   sobrar apenas conteúdo "morto")
  3. Listas OR (packaging)   {Bulk / 1,000 kg Big Bags / 50 kg PP Bags}
                              → substituída pela escolha do usuário
  4. Origem condicional      {BRAZIL: ...} vs {ORIGIN_COUNTRY: ...}  (VegOil)
                              → escolhe um conforme origem
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.document import Document as _Doc
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError as exc:   # pragma: no cover
    raise ImportError(
        "python-docx é obrigatório. Instale com: pip install python-docx"
    ) from exc


# Padrão genérico: qualquer {...} que NÃO contém { ou } internos
BRACE_PATTERN = re.compile(r"\{([^{}]+)\}")


# ═════════════════════════════════════════════════════════════════════════════
# RenderContext
# ═════════════════════════════════════════════════════════════════════════════

@dataclass
class RenderContext:
    """Plano de renderização para uma LOI."""
    # Identificação
    commodity_code: str = ""
    selected_product_label: str = ""

    # Keywords (vindo do dicionário)
    selected_strict_keywords: List[str] = field(default_factory=list)
    selected_family_keywords: List[str] = field(default_factory=list)
    other_strict_keywords:    List[str] = field(default_factory=list)   # de OUTROS produtos

    # Substituição direta de chaves universais
    simple_keys: Dict[str, str] = field(default_factory=dict)

    # Substituições literais (texto SEM chaves {}), ex.: "Date: DD/MM/YYYY"
    literal_subs: Dict[str, str] = field(default_factory=dict)

    # Chaves que, quando resolvidas para "", devem remover a linha INTEIRA do documento
    line_kill_keys: List[str] = field(default_factory=list)

    # Selecao de embalagem (substitui {OptA / OptB / OptC} pelo escolhido)
    selected_packaging: Optional[str] = None

    # Conditional origin (VegOil): "Brazil" → mantém {BRAZIL: ...}
    origin_is_brazil: Optional[bool] = None

    # Marcador de payment terms (selecionado: 'sblc' | 'dlc')
    payment_choice: Optional[str] = None


# ═════════════════════════════════════════════════════════════════════════════
# API pública
# ═════════════════════════════════════════════════════════════════════════════

def render_loi(template_bytes: bytes, ctx: RenderContext) -> bytes:
    """Renderiza LOI segundo o RenderContext."""
    doc = Document(io.BytesIO(template_bytes))

    # 1) Substitui em todos os parágrafos do corpo
    for p in doc.paragraphs:
        _process_paragraph(p, ctx)

    # 2) Em tabelas: por linha, decide manter/dropar antes de processar
    for table in doc.tables:
        _process_table(table, ctx)

    # 3) Cleanup: remove parágrafos do corpo que ficaram totalmente vazios
    _drop_empty_body_paragraphs(doc)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def list_braces_in_template(template_bytes: bytes) -> List[str]:
    """Útil para debug — lista todos os {...} encontrados."""
    doc = Document(io.BytesIO(template_bytes))
    found: set[str] = set()
    for p in doc.paragraphs:
        for m in BRACE_PATTERN.finditer(p.text):
            found.add(m.group(0))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for m in BRACE_PATTERN.finditer(p.text):
                        found.add(m.group(0))
    return sorted(found)


# ═════════════════════════════════════════════════════════════════════════════
# Decisor por padrão {X}
# ═════════════════════════════════════════════════════════════════════════════

def _decide_pattern(inner: str, ctx: RenderContext) -> Optional[str]:
    """
    Decide o que fazer com o conteúdo de um {X}.

    Retorna:
      str  → substitui {X} por essa string
      ""   → remove {X} (drop)
      None → não tocar (deixa {X} no documento)
    """
    # 1) Marcador do produto SELECIONADO → unwrap  (verificar PRIMEIRO para evitar
    #    falsos positivos quando cross-commodity keywords aparecem no texto)
    for kw in ctx.selected_strict_keywords:
        if kw and kw in inner:
            return inner.strip()
    for kw in ctx.selected_family_keywords:
        if kw and kw in inner:
            return inner.strip()

    # 2) Marcador de qualquer OUTRO produto (mesma commodity ou outra) → DROP
    for kw in ctx.other_strict_keywords:
        if kw and kw in inner:
            return ""

    # 3) Origem condicional (VegOil)
    if ctx.origin_is_brazil is not None:
        if inner.startswith("BRAZIL:"):
            if ctx.origin_is_brazil:
                # remove "BRAZIL: " prefix and unwrap
                return inner[len("BRAZIL:"):].strip()
            return ""
        if inner.startswith("ORIGIN_COUNTRY:"):
            if not ctx.origin_is_brazil:
                return inner[len("ORIGIN_COUNTRY:"):].strip()
            return ""

    # 4) Payment terms — texto longo do DLC já tokenizado no template
    if inner.startswith("DLC/LC") or inner.startswith("SBLC"):
        if ctx.payment_choice == "dlc" and inner.startswith("DLC/LC"):
            return inner.strip()
        if ctx.payment_choice == "sblc" and inner.startswith("SBLC"):
            return inner.strip()
        # Se a escolha não corresponde a este padrão, dropa
        return ""

    # 5) OR-list: " / " entre opções → escolher packaging selecionada
    #    ou opção única que bate exatamente com a embalagem escolhida
    if ctx.selected_packaging:
        if " / " in inner:
            opts = [o.strip() for o in inner.split("/")]
            if ctx.selected_packaging in opts:
                return ctx.selected_packaging
            # tentar match insensitive
            low = ctx.selected_packaging.lower()
            for o in opts:
                if o.lower() == low:
                    return o
            # não bateu — conteúdo mantido sem chaves (lista de certifs, etc.)
            return inner.strip()
        # Opção única: match exato → unwrap
        if inner.strip().lower() == ctx.selected_packaging.lower():
            return inner.strip()

    # 6) Chave universal (uppercase identifier)
    # Procura por chave no formato: identificador puro (CITY, YYYY) OU
    # padrões com placeholders (USD XXX.XX, XXX,XXX).
    # Tentamos lookup exato primeiro
    if inner in ctx.simple_keys:
        return ctx.simple_keys[inner]

    # 7) Outros casos (texto manual, etc.) — não tocar
    return None


def _replace_in_text(text: str, ctx: RenderContext) -> str:
    """Aplica decisor a todos os {X} no texto."""
    if "{" not in text:
        return text

    def repl(m: re.Match) -> str:
        inner = m.group(1)
        decision = _decide_pattern(inner, ctx)
        if decision is None:
            # Nenhuma regra específica → remove as chaves mas mantém o conteúdo.
            # Garante que NENHUM {marcador} sobreviva no documento final.
            return inner.strip()
        return decision              # substitui (pode ser "")

    return BRACE_PATTERN.sub(repl, text)


# ═════════════════════════════════════════════════════════════════════════════
# Processamento de parágrafos e tabelas
# ═════════════════════════════════════════════════════════════════════════════

def _clean_destination_separators(text: str) -> str:
    """Remove ·  separators that were left orphaned after empty field substitution.

    Handles patterns like:
      'City ·  · , Country'  →  'City, Country'
      'City · State · , Country'  →  'City · State, Country'
      'City ·  · Port, Country'  →  'City · Port, Country'
    """
    # Two consecutive separators with only whitespace between → single separator
    text = re.sub(r'·\s*·', '·', text)
    # Separator immediately before a comma → just comma
    text = re.sub(r'·\s*,', ',', text)
    # Space(s) immediately before a comma → comma
    text = re.sub(r'\s+,', ',', text)
    # Trailing separator
    text = re.sub(r'\s*·\s*$', '', text.rstrip())
    # Leading separator
    text = re.sub(r'^\s*·\s*', '', text)
    # Leading comma (e.g. when all optional fields are empty and only country remains)
    text = re.sub(r'^\s*,\s*', '', text)
    # Collapse multiple spaces
    text = re.sub(r'  +', ' ', text)
    return text


def _merge_paragraph_runs(paragraph: Paragraph) -> None:
    """Collapse all runs into the first so {markers} split across runs are matchable."""
    runs = paragraph.runs
    if len(runs) <= 1:
        return
    full_text = "".join(r.text for r in runs)
    runs[0].text = full_text
    for r in runs[1:]:
        r.text = ""


def _process_paragraph(paragraph: Paragraph, ctx: RenderContext) -> None:
    """Substitui no parágrafo, preservando o estilo do primeiro run."""
    original_text = paragraph.text
    has_braces  = "{" in original_text
    has_literal = any(lit in original_text for lit in ctx.literal_subs)

    if not has_braces and not has_literal:
        return

    # Line-kill: se o parágrafo contém {KEY} de alguma line_kill_key E essa chave
    # está vazia → remove o parágrafo inteiro antes de qualquer substituição.
    if has_braces and ctx.line_kill_keys:
        for kill_key in ctx.line_kill_keys:
            if f"{{{kill_key}}}" in original_text and not ctx.simple_keys.get(kill_key):
                el = paragraph._element
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
                return

    # Merge runs first so cross-run markers like {SBLC (MT760)...} are matchable
    _merge_paragraph_runs(paragraph)

    runs = paragraph.runs
    current = runs[0].text if runs else paragraph.text

    # 1) Brace substitutions
    if has_braces:
        current = _replace_in_text(current, ctx)

    # 2) Literal substitutions (text without {braces}, e.g. "Date: DD/MM/YYYY")
    for lit, rep in ctx.literal_subs.items():
        if lit in current:
            current = current.replace(lit, rep)

    # 3) Clean up orphan · separators from empty destination fields
    if '·' in current:
        current = _clean_destination_separators(current)

    if runs:
        if current != runs[0].text:
            runs[0].text = current
            for r in runs[1:]:
                r.text = ""
    else:
        paragraph.text = current


def _row_is_dead(row, ctx: RenderContext) -> bool:
    """
    Linha de tabela é "morta" se TODA referência de produto que ela contém
    aponta para produtos NÃO selecionados, e nenhuma menciona o selecionado.

    Verifica tanto marcadores {X} quanto texto literal (templates como CHICKEN
    que usam nomes de produto sem chaves).
    """
    has_selected_ref = False
    has_other_ref    = False
    for cell in row.cells:
        # Texto completo da célula (todas as linhas concatenadas)
        cell_text = " ".join(p.text for p in cell.paragraphs)

        # A) Brace-pattern markers
        for m in BRACE_PATTERN.finditer(cell_text):
            inner = m.group(1)
            for kw in ctx.selected_strict_keywords + ctx.selected_family_keywords:
                if kw and kw in inner:
                    has_selected_ref = True
                    break
            for kw in ctx.other_strict_keywords:
                if kw and kw in inner:
                    has_other_ref = True
                    break

        # B) Plain-text keywords (templates que não usam {markers})
        for kw in ctx.selected_strict_keywords + ctx.selected_family_keywords:
            if kw and kw in cell_text:
                has_selected_ref = True
                break
        for kw in ctx.other_strict_keywords:
            if kw and kw in cell_text:
                has_other_ref = True
                break

    return has_other_ref and not has_selected_ref


def _drop_empty_cell_paragraphs(cell) -> None:
    """Remove parágrafos vazios que sobraram após substituição dentro de uma célula."""
    for p in list(cell.paragraphs):
        if p.text.strip():
            continue
        has_content_run = any(r.text for r in p.runs)
        if not has_content_run:
            continue   # parágrafo vazio natural — manter
        el = p._element
        parent = el.getparent()
        # Não remove se for o único parágrafo da célula (Word exige ≥1)
        if parent is not None and len(list(parent)) > 1:
            parent.remove(el)


def _process_table(table: Table, ctx: RenderContext) -> None:
    """Remove linhas mortas; processa as restantes."""
    rows_to_remove = [row for row in table.rows if _row_is_dead(row, ctx)]
    for row in rows_to_remove:
        el = row._tr
        if el.getparent() is not None:
            el.getparent().remove(el)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _process_paragraph(p, ctx)
            _drop_empty_cell_paragraphs(cell)
            for inner_table in cell.tables:
                _process_table(inner_table, ctx)


def _drop_empty_body_paragraphs(doc: _Doc) -> None:
    """Remove parágrafos do corpo cujo texto ficou completamente vazio."""
    for p in list(doc.paragraphs):
        if p.text.strip() == "":
            # Só remove se foi originalmente um parágrafo com chave (heurística:
            # parágrafo vazio sem runs com texto). Mantém parágrafos vazios
            # que servem de espaçamento se contêm runs com break/space.
            has_content_run = any(r.text for r in p.runs)
            if not has_content_run:
                continue   # parágrafo vazio "natural" — manter
            el = p._element
            if el.getparent() is not None:
                el.getparent().remove(el)


# ═════════════════════════════════════════════════════════════════════════════
# Builder de RenderContext a partir do dicionário + inputs do widget
# ═════════════════════════════════════════════════════════════════════════════

def build_context(
    commodity_code: str,
    product_label: str,
    user_inputs: Dict[str, str],
) -> RenderContext:
    """Monta um RenderContext a partir do dicionário + inputs do widget."""
    from data.knowledge.loi_dictionary import (
        get_commodity, get_product, PAYMENT_TERMS, PERFORMANCE_BOND_TEXTS,
    )
    from datetime import datetime

    com  = get_commodity(commodity_code)
    prod = get_product(commodity_code, product_label)

    # Keywords
    selected_strict = list(prod.get("strict_keywords", []))
    selected_family = list(prod.get("family_keywords", []))
    other_strict: List[str] = []

    # 1) Produtos da mesma commodity (não selecionados)
    for p in com["products"]:
        if p["label"] == product_label:
            continue
        for kw in p.get("strict_keywords", []):
            if kw not in selected_strict and kw not in selected_family and kw not in other_strict:
                other_strict.append(kw)

    # 2) Todas as outras commodities — garante que lixo de cópia de template seja dropado
    from data.knowledge.loi_dictionary import COMMODITIES as _ALL_COMMODITIES
    for _c_code, _c_data in _ALL_COMMODITIES.items():
        if _c_code == commodity_code:
            continue
        for _p in _c_data["products"]:
            for _kw in _p.get("strict_keywords", []):
                if _kw and _kw not in selected_strict and _kw not in selected_family and _kw not in other_strict:
                    other_strict.append(_kw)

    # Simple keys
    today = datetime.now()
    today_str = today.strftime("%d/%m/%Y")
    simple: Dict[str, str] = {
        "YYYY":         str(today.year),
        "YYYYMMDD":     today.strftime("%Y%m%d"),
        "DD/MM/YYYY":   today_str,
    }
    # Inputs do widget (uppercase keys)
    for k, v in user_inputs.items():
        if v is None or str(v).strip() == "":
            continue
        simple[k.upper()] = str(v)

    # Literal substitutions — texto sem {chaves} no template
    # "Date: DD/MM/YYYY" na célula de assinatura não usa {}, precisa de substituição literal
    literal_subs: Dict[str, str] = {
        "Date: DD/MM/YYYY":     f"Date: {today_str}",
        "DATE: DD/MM/YYYY":     f"DATE: {today_str}",
        "date: DD/MM/YYYY":     f"date: {today_str}",
    }

    # ── Chaves derivadas — necessárias para os templates reais ────────────────

    # DESTINATARY: primeiro nome e nome de pessoa (saudação)
    full_name = (
        simple.get("DESTINATARY_LOIFULLNAME")
        or simple.get("DESTINATARY", "")
    ).strip()
    if full_name:
        simple.setdefault("DESTINATARY_LOIFULLNAME",   full_name)
        simple.setdefault("DESTINATARY_LOIPERSON_NAME", full_name)
        first = full_name.split()[0].upper()
        simple.setdefault("DESTINATARY_LOIFIRSTNAME", first)

    # STATE/PROVINCE — templates usam este padrão, widget envia "STATE"
    if "STATE/PROVINCE" not in simple:
        state_val = simple.get("STATE", "")
        if state_val:
            simple["STATE/PROVINCE"] = state_val

    # COUNTRY — templates usam {COUNTRY}, widget pode enviar NATIONALITY_OF_DESTINATION
    if "COUNTRY" not in simple:
        country_val = simple.get("NATIONALITY_OF_DESTINATION", "")
        if country_val:
            simple["COUNTRY"] = country_val

    # Volume + cálculo anual
    monthly = user_inputs.get("VOLUME_MONTHLY")
    if monthly:
        try:
            m_val = int(str(monthly).replace(",", "").replace(".", ""))
            simple["XXX,XXX"]            = f"{m_val:,}".replace(",", ".")
            simple["XXX"]                = f"{m_val:,}".replace(",", ".")
            simple["TOTAL_VOLUME_XXXX,XXX"] = f"{m_val * 12:,}".replace(",", ".")
            simple["TOTAL_VOLUME_XXXX"]    = f"{m_val * 12:,}".replace(",", ".")
        except ValueError:
            pass

    # Target price — quando fornecido substitui {USD XXX.XX}; quando vazio, a
    # linha inteira "Target Price: {USD XXX.XX} per MT." é removida via line_kill.
    line_kill_keys: List[str] = []
    target_price = user_inputs.get("TARGET_PRICE", "")
    if target_price:
        try:
            simple["USD XXX.XX"] = f"{float(str(target_price).replace(',', '.')):.2f}"
        except ValueError:
            simple["USD XXX.XX"] = str(target_price)
    else:
        simple["USD XXX.XX"] = ""
        line_kill_keys.append("USD XXX.XX")

    # Subject manual
    subject = user_inputs.get("SUBJECT")
    if subject:
        simple["TEXTO MANUAL A SER INCERIDO"] = subject

    # Performance bond — sempre registra a chave para evitar que {PERFORMANCE_BOND}
    # caia no fallback e apareça como texto literal. Se o usuário não escolheu
    # nada (ou escolheu Spot/Trial), resolve para "" e o parágrafo é limpo.
    pb = user_inputs.get("PERFORMANCE_BOND", "")
    pb_text = PERFORMANCE_BOND_TEXTS.get(pb, "") if pb else ""
    simple["PERFORMANCE_BOND"] = pb_text

    # Origem (Sugar/VegOil)
    origin = user_inputs.get("ORIGIN_COUNTRY")
    origin_is_brazil = None
    if origin:
        simple["ORIGIN_COUNTRY"] = origin
        origin_is_brazil = (origin.strip().lower() == "brazil")

    # Campo opcional: se não fornecido → "" → parágrafo removido pelo cleanup
    simple.setdefault("TEXTO MANUAL A SER INCERIDO", "")

    # Campos de destino opcionais — sempre em simple_keys (com "" se vazio) para
    # evitar que o fallback retorne o nome da chave literalmente no documento.
    for _opt in ["CITY", "STATE/PROVINCE", "NAME_OF_PORT", "COUNTRY",
                 "NATIONALITY_OF_DESTINATION", "STATE", "ATTN"]:
        simple.setdefault(_opt, "")

    # Cotton: HVI Transparency Clause
    if commodity_code.upper() == "COTTON":
        hvi = com["extra_rules"].get("hvi_transparency_clause", "")
        if hvi:
            simple["HVI_TRANSPARENCY_CLAUSE"] = hvi

    # VegOil: o template usa {BRAZIL: ...} / {ORIGIN_COUNTRY: ...} para Quality
    # Standard (step 3 do decisor). Se o template tiver também {QUALITY_STANDARD}
    # separado, forçamos "" para evitar duplicação.
    if commodity_code.upper() == "VEGOIL":
        simple["QUALITY_STANDARD"] = ""

    # Payment choice (sblc/dlc)
    payment_in = user_inputs.get("PAYMENT_TERMS", "")
    payment_choice = None
    if "DLC" in payment_in.upper():
        payment_choice = "dlc"
    elif "SBLC" in payment_in.upper():
        payment_choice = "sblc"

    return RenderContext(
        commodity_code=commodity_code,
        selected_product_label=product_label,
        selected_strict_keywords=selected_strict,
        selected_family_keywords=selected_family,
        other_strict_keywords=other_strict,
        simple_keys=simple,
        literal_subs=literal_subs,
        line_kill_keys=line_kill_keys,
        selected_packaging=user_inputs.get("PACKAGING"),
        origin_is_brazil=origin_is_brazil,
        payment_choice=payment_choice,
    )


# ═════════════════════════════════════════════════════════════════════════════
# Filename builder
# ═════════════════════════════════════════════════════════════════════════════

def build_output_filename(
    destinatary_first_name: str,
    commodity_code: str,
    today_yyyymmdd: Optional[str] = None,
) -> str:
    """LOI-SE-{YYYYMMDD}-{NAME}-{COMMODITY_EN}-{YYYY}.docx"""
    from datetime import datetime
    from data.knowledge.loi_dictionary import get_commodity

    com = get_commodity(commodity_code)
    today = today_yyyymmdd or datetime.now().strftime("%Y%m%d")
    year  = today[:4]
    name  = destinatary_first_name.strip().upper().replace(" ", "_")
    com_en = com["label_en"].upper().replace(" ", "_")
    return f"LOI-SE-{today}-{name}-{com_en}-{year}.docx"
