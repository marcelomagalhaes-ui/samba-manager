# -*- coding: utf-8 -*-
"""
services/imfpa_template_engine.py
==================================
Engine de renderização de IMFPA.

Responsabilidades:
  1. Substituição direta de marcadores {MARKER} → valor
  2. Tratamento de marcadores fragmentados em múltiplos runs
  3. Substituição do literal SOYBEAN → commodity escolhida
  4. Substituição dos literais "USD X.XX per MT" (taxa total e por embarque)
  5. Iteração em parágrafos, tabelas, cabeçalho, rodapé

Marcadores suportados (N = 1, 2 ou 3 conforme número de partes):
  Por parte N
  -----------
  {FULL_NAME_Company_N}            Nome completo da empresa N
  {COUNTRY_N}                      País da empresa N
  {TAX_ID_Company_N}               CNPJ/EIN/VAT da empresa N
  {REGISTERED_ADDRESS_Company_N}   Endereço registrado
  {FULL_NAME_PERSON_N}             Nome do representante legal
  {PASSPORT_CODE_PERSON_N}         Número de passaporte
  {BENEFICIARY_NAME_COMPANY_N}     Nome do beneficiário bancário
  {DOCUMENT_NUMBER_ COMPANY_N}     (com espaço antes de COMPANY) Número do doc bancário
  {BANK NAME_ COMPANY_N}           (com espaço antes de COMPANY) Nome do banco
  {SWIFT_CODE_ COMPANY_N}          SWIFT
  {IBAN_CODE_ COMPANY_N}           IBAN

  Transação
  ---------
  {DD / MM / YYYY}  ou  {DD/MM/YYYY}   Data formatada
  {DDMMYY-FIRST_NAME_COMPANY_1}        Código do documento (ex: 140526-SAMBA)
  {DDMMYY-FIRSTNAMECOMPANY2}           Código do documento (variante)
  {000.000}                            Quantidade em MT (ex: 15,000)
  {SPA_CODE}                           Código da SPA

  Literais a substituir
  ---------------------
  SOYBEAN          → nome da commodity (ex: SOYBEAN MEAL, CORN, SUGAR etc.)
  USD X.XX per MT  → "USD {fee_per_mt} per MT"  (primeira ocorrência = por embarque,
                     segunda ocorrência = taxa total — ou vice-versa)
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError as exc:  # pragma: no cover
    raise ImportError("python-docx é obrigatório. Instale com: pip install python-docx") from exc

# ─── Padrões ─────────────────────────────────────────────────────────────────
# Marcador genérico {…} sem chaves internas
_BRACE = re.compile(r"\{([^{}]+)\}")

# Literais fixos
_SOYBEAN_PATTERN  = re.compile(r"\bSOYBEAN\b")
_USD_PATTERN      = re.compile(r"USD\s+X\.XX\s+per\s+MT", re.IGNORECASE)


# ═════════════════════════════════════════════════════════════════════════════
# Contexto
# ═════════════════════════════════════════════════════════════════════════════

@dataclass
class IMFPAContext:
    """
    Todos os valores necessários para preencher um template IMFPA.

    Campos de parte: dicionários indexados por N (1, 2, 3).
    """

    # --- Transação ---
    date_str:      str = ""   # "14/05/2026"
    doc_code:      str = ""   # "140526-SAMBA"  (parte depois do traço)
    quantity_mt:   str = ""   # "15,000"
    spa_code:      str = ""   # "SPA-2026-001"
    commodity:     str = ""   # "SOYBEAN MEAL"

    # Taxas: fee_per_shipment e fee_total são substituídas nessa ordem
    # nas duas ocorrências de "USD X.XX per MT"
    fee_per_shipment: str = ""   # ex: "2.50"
    fee_total:        str = ""   # ex: "2.50"  (pode ser igual)

    # --- Por parte: N = 1, 2, 3 ---
    company_name:        Dict[int, str] = field(default_factory=dict)
    country:             Dict[int, str] = field(default_factory=dict)
    tax_id:              Dict[int, str] = field(default_factory=dict)
    address:             Dict[int, str] = field(default_factory=dict)
    legal_rep_name:      Dict[int, str] = field(default_factory=dict)
    passport:            Dict[int, str] = field(default_factory=dict)
    beneficiary_name:    Dict[int, str] = field(default_factory=dict)
    doc_number:          Dict[int, str] = field(default_factory=dict)
    bank_name:           Dict[int, str] = field(default_factory=dict)
    swift:               Dict[int, str] = field(default_factory=dict)
    iban:                Dict[int, str] = field(default_factory=dict)

    def flat_replacements(self) -> Dict[str, str]:
        """
        Retorna dicionário {marcador_sem_chaves: valor} para substituição direta.
        Inclui variantes com/sem espaço, com/sem barra nos marcadores de data.
        """
        r: Dict[str, str] = {}

        # Data (duas variantes)
        if self.date_str:
            r["DD / MM / YYYY"] = self.date_str
            r["DD/MM/YYYY"]     = self.date_str

        # Código do documento (duas variantes de marcador)
        # O template usa {DDMMYY-FIRST_NAME_COMPANY_1} e {DDMMYY-FIRSTNAMECOMPANY2}
        # Ambos recebem o mesmo doc_code construído pelo chamador
        if self.doc_code:
            r["DDMMYY-FIRST_NAME_COMPANY_1"] = self.doc_code
            r["DDMMYY-FIRSTNAMECOMPANY2"]    = self.doc_code

        if self.quantity_mt:
            r["000.000"] = self.quantity_mt

        if self.spa_code:
            r["SPA_CODE"] = self.spa_code

        # Por parte
        for n in (1, 2, 3):
            if self.company_name.get(n):
                r[f"FULL_NAME_Company_{n}"] = self.company_name[n]
            if self.country.get(n):
                r[f"COUNTRY_{n}"] = self.country[n]
            if self.tax_id.get(n):
                r[f"TAX_ID_Company_{n}"] = self.tax_id[n]
            if self.address.get(n):
                r[f"REGISTERED_ADDRESS_Company_{n}"] = self.address[n]
            if self.legal_rep_name.get(n):
                r[f"FULL_NAME_PERSON_{n}"] = self.legal_rep_name[n]
            if self.passport.get(n):
                r[f"PASSPORT_CODE_PERSON_{n}"] = self.passport[n]
            if self.beneficiary_name.get(n):
                r[f"BENEFICIARY_NAME_COMPANY_{n}"] = self.beneficiary_name[n]
            # Marcadores com espaço antes de COMPANY (como estão nos templates)
            if self.doc_number.get(n):
                r[f"DOCUMENT_NUMBER_ COMPANY_{n}"] = self.doc_number[n]
            if self.bank_name.get(n):
                r[f"BANK NAME_ COMPANY_{n}"] = self.bank_name[n]
            if self.swift.get(n):
                r[f"SWIFT_CODE_ COMPANY_{n}"] = self.swift[n]
            if self.iban.get(n):
                r[f"IBAN_CODE_ COMPANY_{n}"] = self.iban[n]

        return r


# ═════════════════════════════════════════════════════════════════════════════
# Helpers de substituição em runs
# ═════════════════════════════════════════════════════════════════════════════

def _merge_paragraph_text(para) -> str:
    """Retorna o texto completo do parágrafo (concatenação dos runs)."""
    return "".join(r.text for r in para.runs)


def _set_paragraph_runs(para, new_text: str) -> None:
    """
    Substitui o conteúdo textual do parágrafo preservando formatação do primeiro run.
    Estratégia: coloca todo o texto no primeiro run e zera os demais.
    """
    if not para.runs:
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""


def _replace_in_paragraph(para, replacements: Dict[str, str],
                           commodity: str,
                           usd_state: List[int]) -> None:
    """
    Aplica todas as substituições em um parágrafo:
      1. Reconstrói texto completo a partir dos runs
      2. Aplica {MARKER} → valor
      3. Substitui SOYBEAN → commodity
      4. Substitui USD X.XX per MT (com contador de ocorrências via usd_state)
      5. Reescreve os runs
    """
    full = _merge_paragraph_text(para)
    if not full.strip():
        return

    changed = False

    # 1. Marcadores {…}
    def _sub_marker(m: re.Match) -> str:
        key = m.group(1)
        if key in replacements:
            return replacements[key]
        return m.group(0)   # mantém intacto se não mapeado

    new_text, n_subs = _BRACE.subn(_sub_marker, full)
    if n_subs:
        changed = True
        full = new_text

    # 2. SOYBEAN literal
    if commodity and _SOYBEAN_PATTERN.search(full):
        full = _SOYBEAN_PATTERN.sub(commodity, full)
        changed = True

    # 3. USD X.XX per MT (duas ocorrências → per_shipment, total)
    def _sub_usd(m: re.Match) -> str:
        idx = usd_state[0]
        usd_state[0] += 1
        # A primeira ocorrência é por embarque, a segunda é o total
        # (ajuste conforme template real — pode ser invertido)
        if idx == 0:
            return f"USD {usd_state[1]} per MT"   # per_shipment armazenado em [1]
        else:
            return f"USD {usd_state[2]} per MT"   # total armazenado em [2]

    # Usamos um closure para capturar os valores reais fora do padrão state
    if _USD_PATTERN.search(full):
        # usd_state = [count, per_shipment_str, total_str]
        def _sub_usd_inner(m: re.Match) -> str:
            cnt = usd_state[0]
            usd_state[0] += 1
            val = usd_state[1] if cnt == 0 else usd_state[2]
            return f"USD {val} per MT"

        full = _USD_PATTERN.sub(_sub_usd_inner, full)
        changed = True

    if changed:
        _set_paragraph_runs(para, full)


def _iter_paragraphs(doc: "_Doc"):
    """Itera por todos os parágrafos do documento: corpo, tabelas, headers, footers."""
    # Corpo
    for para in doc.paragraphs:
        yield para

    # Tabelas (aninhadas inclusive)
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)

    # Cabeçalho / rodapé de todas as seções
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr is not None:
                for para in hdr.paragraphs:
                    yield para
                for table in hdr.tables:
                    yield from _iter_table_paragraphs(table)
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            if ftr is not None:
                for para in ftr.paragraphs:
                    yield para
                for table in ftr.tables:
                    yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    """Itera por parágrafos de uma tabela (suporta tabelas aninhadas)."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                yield para
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


# ═════════════════════════════════════════════════════════════════════════════
# API pública
# ═════════════════════════════════════════════════════════════════════════════

def render_imfpa(template_bytes: bytes, ctx: IMFPAContext) -> bytes:
    """
    Recebe os bytes de um .docx de template IMFPA e um contexto preenchido.
    Retorna os bytes do .docx renderizado.
    """
    doc = Document(io.BytesIO(template_bytes))

    replacements = ctx.flat_replacements()
    # usd_state: [count, per_shipment_value, total_value]
    usd_state = [0, ctx.fee_per_shipment or "X.XX", ctx.fee_total or "X.XX"]

    for para in _iter_paragraphs(doc):
        _replace_in_paragraph(para, replacements, ctx.commodity, usd_state)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_imfpa_output_filename(doc_code: str, n_parties: int) -> str:
    """
    Gera o nome do arquivo de saída.
    Ex: "140526-SAMBA_IMFPA_2parties.docx"
    """
    safe_code = (doc_code or "IMFPA").replace("/", "-").replace(" ", "_")
    return f"{safe_code}_IMFPA_{n_parties}parties.docx"
