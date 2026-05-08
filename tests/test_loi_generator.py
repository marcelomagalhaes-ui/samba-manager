# -*- coding: utf-8 -*-
"""
tests/test_loi_generator.py
============================
Testes unitários do LOI Generator (engine + dicionário + agente).

Não dependem do Drive — usam templates DOCX gerados em memória.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from data.knowledge.loi_dictionary import (
    COMMODITIES, get_commodity, get_product, list_product_labels,
)
from services.loi_template_engine import (
    RenderContext,
    build_context,
    build_output_filename,
    list_braces_in_template,
    render_loi,
)


# ═════════════════════════════════════════════════════════════════════════════
# Fixtures
# ═════════════════════════════════════════════════════════════════════════════

def _make_template(lines: list[str]) -> bytes:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _make_template_with_table(lines: list[str], table_rows: list[list[str]]) -> bytes:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, val in enumerate(row):
                t.cell(i, j).text = val
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _doc_text(output_bytes: bytes) -> str:
    doc = Document(io.BytesIO(output_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


# ═════════════════════════════════════════════════════════════════════════════
# Dicionário
# ═════════════════════════════════════════════════════════════════════════════

def test_dictionary_has_all_commodities():
    expected = {"SOY", "CORN", "SUGAR", "CHICKEN", "VEGOIL", "COTTON", "RICE"}
    assert set(COMMODITIES.keys()) == expected


def test_dictionary_structure():
    for code, com in COMMODITIES.items():
        assert "products" in com and com["products"], f"{code}: missing products"
        assert "template_filename" in com
        assert com["template_filename"].startswith("MODEL-LOI-")
        assert "incoterms" in com and com["incoterms"]
        for p in com["products"]:
            assert "label" in p
            assert "strict_keywords" in p
            assert "family_keywords" in p
            assert "packaging_options" in p


def test_get_product_soy():
    p = get_product("SOY", "Yellow Soybean, GMO - Grade #2")
    assert "Yellow Soybean" in p["strict_keywords"]
    assert "Bulk" in p["packaging_options"]


def test_get_product_raises_unknown():
    with pytest.raises(KeyError):
        get_product("SOY", "NonExistentProduct")


def test_list_product_labels_soy():
    labels = list_product_labels("SOY")
    assert "Yellow Soybean, GMO - Grade #2" in labels
    assert "Soybean Meal - Solvent Extracted" in labels


def test_cotton_family_keywords():
    p = get_product("COTTON", "Cotton Lint - 31-1 Good Middling")
    assert "Cotton Lint" in p["family_keywords"]
    assert "31-1" in p["strict_keywords"]


def test_rice_has_five_products():
    assert len(list_product_labels("RICE")) == 5


def test_vegoil_crude_palm_keywords():
    p = get_product("VEGOIL", "Crude Palm Oil (CPO)")
    # CPO tem múltiplas variações de keyword para pegar todas as ocorrências no template
    assert any("CPO" in kw for kw in p["strict_keywords"])


# ═════════════════════════════════════════════════════════════════════════════
# RenderContext — build_context
# ═════════════════════════════════════════════════════════════════════════════

def test_build_context_soy_basic():
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "John Carter",
        "VOLUME_MONTHLY": "50000",
    })
    assert ctx.commodity_code == "SOY"
    assert ctx.selected_product_label == "Yellow Soybean, GMO - Grade #2"
    assert "Yellow Soybean" in ctx.selected_strict_keywords
    assert "Soybean Meal" in ctx.other_strict_keywords
    assert ctx.simple_keys.get("DESTINATARY_LOIFULLNAME") == "John Carter"
    assert "XXX,XXX" in ctx.simple_keys
    assert "TOTAL_VOLUME_XXXX,XXX" in ctx.simple_keys
    assert ctx.simple_keys.get("YYYY") is not None


def test_build_context_payment_sblc():
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
        "PAYMENT_TERMS": "SBLC (MT760)",
    })
    assert ctx.payment_choice == "sblc"


def test_build_context_payment_dlc():
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
        "PAYMENT_TERMS": "DLC (MT700)",
    })
    assert ctx.payment_choice == "dlc"


def test_build_context_vegoil_brazil():
    ctx = build_context("VEGOIL", "Crude Soybean Oil (CDSO)", {
        "DESTINATARY_LOIFULLNAME": "Test",
        "ORIGIN_COUNTRY": "Brazil",
    })
    assert ctx.origin_is_brazil is True
    # QUALITY_STANDARD is "" for VEGOIL — the {BRAZIL:...}/{ORIGIN_COUNTRY:...}
    # conditionals in the template handle quality standard exclusively.
    assert ctx.simple_keys.get("QUALITY_STANDARD") == ""


def test_build_context_vegoil_non_brazil():
    ctx = build_context("VEGOIL", "Crude Soybean Oil (CDSO)", {
        "DESTINATARY_LOIFULLNAME": "Test",
        "ORIGIN_COUNTRY": "Ukraine",
    })
    assert ctx.origin_is_brazil is False
    assert ctx.simple_keys.get("QUALITY_STANDARD") == ""


def test_build_context_performance_bond():
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
        "PERFORMANCE_BOND": "12 months",
    })
    pb = ctx.simple_keys.get("PERFORMANCE_BOND", "")
    assert "2%" in pb and "Seller" in pb


def test_build_context_performance_bond_spot():
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
        "PERFORMANCE_BOND": "Spot/Trial",
    })
    # Spot/Trial → texto vazio → linha será removida pelo engine
    assert ctx.simple_keys.get("PERFORMANCE_BOND", "___") == ""


def test_build_context_target_price():
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
        "TARGET_PRICE": "320.5",
    })
    assert ctx.simple_keys.get("USD XXX.XX") == "USD 320.50"


# ═════════════════════════════════════════════════════════════════════════════
# render_loi — substituição simples
# ═════════════════════════════════════════════════════════════════════════════

def test_render_simple_keys():
    tpl = _make_template([
        "City: {CITY}",
        "Year: {YYYY}",
        "Dest: {DESTINATARY_LOIFULLNAME}",
    ])
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "Maria Silva",
        "CITY": "Santos",
    })
    out = render_loi(tpl, ctx)
    text = _doc_text(out)
    assert "Maria Silva" in text
    assert "Santos" in text
    assert "{CITY}" not in text
    assert "{DESTINATARY_LOIFULLNAME}" not in text


def test_render_product_marker_kept():
    """Marcador do produto SELECIONADO é mantido (braces removidas)."""
    tpl = _make_template([
        "Product: {Yellow Soybean, GMO - Grade #2}",
    ])
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
    })
    out = render_loi(tpl, ctx)
    text = _doc_text(out)
    assert "Yellow Soybean, GMO - Grade #2" in text
    assert "{Yellow Soybean" not in text


def test_render_other_product_dropped():
    """Marcadores de OUTROS produtos são removidos."""
    tpl = _make_template([
        "A: {Yellow Soybean}",
        "B: {Soybean Meal}",
    ])
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
    })
    out = render_loi(tpl, ctx)
    text = _doc_text(out)
    assert "Yellow Soybean" in text
    assert "Soybean Meal" not in text


def test_render_packaging_or_list():
    """OR-list de packaging substituída pela opção selecionada."""
    tpl = _make_template([
        "Packaging: {Bulk / 1,000 kg Big Bags / 50 kg PP Bags}",
    ])
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
        "PACKAGING": "Bulk",
    })
    out = render_loi(tpl, ctx)
    text = _doc_text(out)
    assert "Packaging: Bulk" in text
    assert "Big Bags" not in text


def test_render_empty_paragraph_removed():
    """Parágrafo cujo único conteúdo era uma chave substituída por '' deve sumir."""
    tpl = _make_template([
        "Before",
        "{PERFORMANCE_BOND}",
        "After",
    ])
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
        "PERFORMANCE_BOND": "Spot/Trial",   # → texto vazio
    })
    out = render_loi(tpl, ctx)
    text = _doc_text(out)
    assert "Before" in text
    assert "After" in text
    assert "{PERFORMANCE_BOND}" not in text


# ═════════════════════════════════════════════════════════════════════════════
# render_loi — tabelas
# ═════════════════════════════════════════════════════════════════════════════

def test_render_table_row_dead_removed():
    """Linha de tabela com referência apenas a produto não selecionado é removida."""
    tpl = _make_template_with_table(
        lines=["Header"],
        table_rows=[
            ["Yellow Soybean spec", "{Yellow Soybean}"],
            ["Soybean Meal spec",   "{Soybean Meal}"],
        ],
    )
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
    })
    out = render_loi(tpl, ctx)
    doc = Document(io.BytesIO(out))
    table = doc.tables[0]
    assert len(table.rows) == 1
    cells = [c.text for r in table.rows for c in r.cells]
    assert any("Yellow Soybean" in c for c in cells)
    assert all("Soybean Meal" not in c for c in cells)


def test_render_table_row_with_selected_kept():
    """Linha com referência ao produto selecionado é mantida."""
    tpl = _make_template_with_table(
        lines=["Header"],
        table_rows=[
            ["A", "{Yellow Soybean}"],
            ["B", "{Soybean Meal}"],
            ["C", "Static row"],
        ],
    )
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
    })
    out = render_loi(tpl, ctx)
    doc = Document(io.BytesIO(out))
    table = doc.tables[0]
    # Row B (Soybean Meal) removed; rows A and C kept
    assert len(table.rows) == 2


# ═════════════════════════════════════════════════════════════════════════════
# VegOil conditional origin
# ═════════════════════════════════════════════════════════════════════════════

def test_render_vegoil_brazil_block():
    tpl = _make_template([
        "Quality: {BRAZIL: FOSFA 54 / ANEC 81}",
        "Quality: {ORIGIN_COUNTRY: FOSFA 54}",
    ])
    ctx = build_context("VEGOIL", "Crude Soybean Oil (CDSO)", {
        "DESTINATARY_LOIFULLNAME": "X",
        "ORIGIN_COUNTRY": "Brazil",
    })
    out = render_loi(tpl, ctx)
    text = _doc_text(out)
    assert "ANEC 81" in text
    assert "ORIGIN_COUNTRY:" not in text
    assert "BRAZIL:" not in text


def test_render_vegoil_non_brazil_block():
    tpl = _make_template([
        "Quality: {BRAZIL: FOSFA 54 / ANEC 81}",
        "Quality: {ORIGIN_COUNTRY: FOSFA 54 only}",
    ])
    ctx = build_context("VEGOIL", "Crude Soybean Oil (CDSO)", {
        "DESTINATARY_LOIFULLNAME": "X",
        "ORIGIN_COUNTRY": "Ukraine",
    })
    out = render_loi(tpl, ctx)
    text = _doc_text(out)
    assert "FOSFA 54 only" in text
    assert "ANEC 81" not in text


# ═════════════════════════════════════════════════════════════════════════════
# Payment terms
# ═════════════════════════════════════════════════════════════════════════════

def test_render_payment_sblc_kept():
    tpl = _make_template([
        "Terms: {SBLC (MT760) Irrevocable, Transferable, Cash Backed, Issued by a Prime Bank}",
        "Terms: {DLC/LC (MT700) Irrevocable, Transferable, Cash Backed, Issued by a Prime Bank}",
    ])
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
        "PAYMENT_TERMS": "SBLC (MT760)",
    })
    out = render_loi(tpl, ctx)
    text = _doc_text(out)
    assert "SBLC" in text
    assert "DLC/LC" not in text


def test_render_payment_dlc_kept():
    tpl = _make_template([
        "Terms: {SBLC (MT760) something}",
        "Terms: {DLC/LC (MT700) something}",
    ])
    ctx = build_context("SOY", "Yellow Soybean, GMO - Grade #2", {
        "DESTINATARY_LOIFULLNAME": "X",
        "PAYMENT_TERMS": "DLC (MT700)",
    })
    out = render_loi(tpl, ctx)
    text = _doc_text(out)
    assert "DLC/LC" in text
    assert "SBLC" not in text


# ═════════════════════════════════════════════════════════════════════════════
# list_braces_in_template
# ═════════════════════════════════════════════════════════════════════════════

def test_list_braces():
    tpl = _make_template(["{CITY}", "{YYYY}", "Plain text", "{Yellow Soybean}"])
    braces = list_braces_in_template(tpl)
    assert "{CITY}" in braces
    assert "{YYYY}" in braces
    assert "{Yellow Soybean}" in braces
    assert "Plain text" not in braces


# ═════════════════════════════════════════════════════════════════════════════
# build_output_filename
# ═════════════════════════════════════════════════════════════════════════════

def test_output_filename_soy():
    fn = build_output_filename("John", "SOY", today_yyyymmdd="20260428")
    assert fn == "LOI-SE-20260428-JOHN-SOYBEAN-2026.docx"


def test_output_filename_vegoil():
    fn = build_output_filename("Maria", "VEGOIL", today_yyyymmdd="20260501")
    assert fn == "LOI-SE-20260501-MARIA-VEGETABLE_OILS-2026.docx"


def test_output_filename_rice():
    fn = build_output_filename("Carlos", "RICE", today_yyyymmdd="20260615")
    assert fn == "LOI-SE-20260615-CARLOS-RICE-2026.docx"


# ═════════════════════════════════════════════════════════════════════════════
# LOIGeneratorAgent — dry run end-to-end (mock template)
# ═════════════════════════════════════════════════════════════════════════════

def test_agent_dry_run_soy(monkeypatch):
    from agents.loi_generator_agent import LOIGeneratorAgent

    fake_template = _make_template([
        "Dear {DESTINATARY_LOIFULLNAME},",
        "City: {CITY}",
        "Product: {Yellow Soybean, GMO - Grade #2}",
        "Meal ref: {Soybean Meal}",
        "Packaging: {Bulk / 1,000 kg Big Bags / 50 kg PP Bags}",
        "Payment: {SBLC (MT760) Irrevocable, Cash Backed}",
        "Payment: {DLC/LC (MT700) Irrevocable, Cash Backed}",
        "Bond: {PERFORMANCE_BOND}",
        "Year: {YYYY}",
    ])

    class FakeDrive:
        def find_file_by_name(self, name, folder_id, ignore_underscore_prefix=True):
            return {"id": "fake-id", "name": name,
                    "mimeType": "application/vnd.google-apps.document"}
        def fetch_as_docx_bytes(self, meta):
            return fake_template
        def upload_file_bytes(self, filename, content, folder_id,
                              mime_type=None, save_as_google_doc=True):
            return {"id": "out-id", "name": filename,
                    "webViewLink": "http://x",
                    "mimeType": "application/vnd.google-apps.document"}

    agent = LOIGeneratorAgent(drive=FakeDrive())
    result = agent.run({
        "commodity_code": "SOY",
        "product_label":  "Yellow Soybean, GMO - Grade #2",
        "user_inputs": {
            "DESTINATARY_LOIFULLNAME": "John Carter",
            "CITY": "Veracruz",
            "VOLUME_MONTHLY": "50000",
            "PACKAGING": "Bulk",
            "PAYMENT_TERMS": "SBLC (MT760)",
            "PERFORMANCE_BOND": "12 months",
        },
        "dry_run": True,
    })
    assert result["status"] == "success"
    assert result["dry_run"] is True
    assert result["filename"].startswith("LOI-SE-")
    assert "SOYBEAN" in result["filename"]
    # Yellow Soybean é selecionado → in selected_keywords
    all_kw = result["selected_keywords"]
    assert "Yellow Soybean" in all_kw
    # Soybean Meal é droppado → in drop_keywords
    assert "Soybean Meal" in result["drop_keywords"]


def test_agent_dry_run_chicken(monkeypatch):
    from agents.loi_generator_agent import LOIGeneratorAgent

    fake_template = _make_template([
        "Dear {DESTINATARY_LOIFULLNAME},",
        "Product: {Chicken Paws} ref",
        "Product: {Chicken Feet} ref",
        "Product: {Chicken Breast} ref",
    ])

    class FakeDrive:
        def find_file_by_name(self, name, folder_id, ignore_underscore_prefix=True):
            return {"id": "fake-id", "name": name,
                    "mimeType": "application/vnd.google-apps.document"}
        def fetch_as_docx_bytes(self, meta):
            return fake_template
        def upload_file_bytes(self, **kw):
            return {}

    agent = LOIGeneratorAgent(drive=FakeDrive())
    result = agent.run({
        "commodity_code": "CHICKEN",
        "product_label":  "Frozen Chicken Paws",
        "user_inputs": {
            "DESTINATARY_LOIFULLNAME": "Ibrahim Al-Rashid",
            "PACKAGING": "40HQ Reefer Containers",
        },
        "dry_run": True,
    })
    assert result["status"] == "success"
    assert "Chicken Paws" in result["selected_keywords"]
    assert "Chicken Feet" in result["drop_keywords"]
    assert "Chicken Breast" in result["drop_keywords"]


def test_agent_rejects_missing_destinatary():
    from agents.loi_generator_agent import LOIGeneratorAgent

    agent = LOIGeneratorAgent(drive=object())
    result = agent.run({
        "commodity_code": "SOY",
        "product_label":  "Yellow Soybean, GMO - Grade #2",
        "user_inputs": {},
    })
    assert result["status"] == "error"
    assert "DESTINATARY" in result["error"]


def test_agent_rejects_unknown_commodity():
    from agents.loi_generator_agent import LOIGeneratorAgent

    agent = LOIGeneratorAgent(drive=object())
    result = agent.run({
        "commodity_code": "XPTO",
        "product_label":  "anything",
        "user_inputs": {"DESTINATARY_LOIFULLNAME": "X"},
    })
    assert result["status"] == "error"
    assert "Commodity desconhecida" in result["error"]


def test_agent_rejects_invalid_product():
    from agents.loi_generator_agent import LOIGeneratorAgent

    agent = LOIGeneratorAgent(drive=object())
    result = agent.run({
        "commodity_code": "SOY",
        "product_label":  "Produto Inexistente",
        "user_inputs": {"DESTINATARY_LOIFULLNAME": "X"},
    })
    assert result["status"] == "error"
    assert "Produto inválido" in result["error"]


def test_agent_sugar_non_brazil_alert(monkeypatch):
    from agents.loi_generator_agent import LOIGeneratorAgent

    fake_template = _make_template([
        "Sugar: {ICUMSA 45} ref",
        "Sugar: {VHP} ref",
    ])

    class FakeDrive:
        def find_file_by_name(self, name, folder_id, ignore_underscore_prefix=True):
            return {"id": "x", "name": name,
                    "mimeType": "application/vnd.google-apps.document"}
        def fetch_as_docx_bytes(self, meta):
            return fake_template

    agent = LOIGeneratorAgent(drive=FakeDrive())
    result = agent.run({
        "commodity_code": "SUGAR",
        "product_label":  "Refined White Sugar ICUMSA 45 - Grade A",
        "user_inputs": {
            "DESTINATARY_LOIFULLNAME": "Ahmad",
            "ORIGIN_COUNTRY": "Thailand",
        },
        "dry_run": True,
    })
    assert result["status"] == "success"
    assert any("MAPA" in a for a in result["alerts"])


def test_agent_full_upload(monkeypatch):
    """Smoke test do fluxo completo (sem dry_run) com mock Drive."""
    from agents.loi_generator_agent import LOIGeneratorAgent

    fake_template = _make_template([
        "Dear {DESTINATARY_LOIFULLNAME},",
        "Year: {YYYY}",
    ])
    uploaded = {}

    class FakeDrive:
        def find_file_by_name(self, name, folder_id, ignore_underscore_prefix=True):
            return {"id": "tmpl-id", "name": name,
                    "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        def fetch_as_docx_bytes(self, meta):
            return fake_template
        def upload_file_bytes(self, filename, content, folder_id,
                              mime_type=None, save_as_google_doc=True):
            uploaded["filename"] = filename
            uploaded["size"] = len(content)
            return {"id": "out-123", "name": filename,
                    "webViewLink": "http://drive/out-123",
                    "mimeType": "application/pdf"}
        def export_gdoc_as_pdf_bytes(self, file_id):
            return b"%PDF-fake"
        def delete_file(self, file_id):
            return True

    agent = LOIGeneratorAgent(drive=FakeDrive())
    result = agent.run({
        "commodity_code": "RICE",
        "product_label":  "White Rice (Polished) - Type 1",
        "output_format":  "pdf",
        "user_inputs": {
            "DESTINATARY_LOIFULLNAME": "Nguyen Van A",
        },
    })
    assert result["status"] == "success"
    assert result["file_id"] == "out-123"
    assert result["web_link"] == "http://drive/out-123"
    assert uploaded["filename"].startswith("LOI-SE-")
    assert "RICE" in uploaded["filename"]
