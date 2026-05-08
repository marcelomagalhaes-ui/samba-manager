# -*- coding: utf-8 -*-
"""
agents/ncnda_generator_agent.py
================================
NCNDAGeneratorAgent — gera o documento NCNDA (EN ou PTBR) com número
dinâmico de partes (Party I = Samba Export + até 3 adicionais).

Pipeline (com suporte a templates Drive):
  1. Recebe payload {language, template_prefix, parties, dry_run}
  2. Tenta localizar template no Drive (ex.: "1NCNDA - SAMBA INTERM DE NEGOCIOS (EN)")
     → Se encontrado: baixa, substitui placeholders {{P2_NAME}} etc., usa o docx
     → Se não encontrado: gera programaticamente (fallback)
  3. Upload no Drive como Google Doc → exporta como PDF
  4. Retorna {status, filename, web_link, size_bytes, file_bytes, template_used, alerts}

Matrix de templates (pasta TEMPLATES_FOLDER_ID):
  EN   + 1 parte → "1NCNDA - SAMBA INTERM DE NEGOCIOS (EN)"
  EN   + 2 partes → "2NCNDA - SAMBA INTERM DE NEGOCIOS (EN)"
  EN   + 3 partes → "3NCNDA - SAMBA INTERM DE NEGOCIOS (EN)"
  PTBR + 1 parte → "1PTNCNDA - SAMBA INTERM DE NEGOCIOS (PT)"
  PTBR + 2 partes → "2PTNCNDA - SAMBA INTERM DE NEGOCIOS (PT)"
  PTBR + 3 partes → "3PTNCNDA - SAMBA INTERM DE NEGOCIOS (PT)"

Placeholders esperados nos templates:
  {{EXEC_DATE}}, {{REF_CODE}},
  {{P2_NAME}}, {{P2_SHORTNAME}}, {{P2_COUNTRY}}, {{P2_TAXID}},
  {{P2_ADDRESS}}, {{P2_REP}}, {{P2_PASSPORT}}
  (idem P3_ e P4_ para templates de 2 e 3 partes)
"""
from __future__ import annotations

import io
import datetime
import sys
from pathlib import Path
from typing import Any, List, Optional

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from agents.base_agent import BaseAgent
from services.google_drive import DriveManager

# python-docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Configurações ─────────────────────────────────────────────────────────────
OUTPUT_FOLDER_ID    = "1CPlF_9TtEZ32B5eTAb4b4jRC-h7L10Z2"
TEMPLATES_FOLDER_ID = "1QMQtu13QaPW9WKa-A5Pv9cKRSus2xVlu"

# Mapeamento (language, n_parties_adicionais) → nome do template no Drive
TEMPLATE_MAP: dict[tuple, str] = {
    ("EN",   1): "1NCNDA - SAMBA INTERM DE NEGOCIOS (EN)",
    ("EN",   2): "2NCNDA - SAMBA INTERM DE NEGOCIOS (EN)",
    ("EN",   3): "3NCNDA - SAMBA INTERM DE NEGOCIOS (EN)",
    ("PTBR", 1): "1PTNCNDA - SAMBA INTERM DE NEGOCIOS (PT)",
    ("PTBR", 2): "2PTNCNDA - SAMBA INTERM DE NEGOCIOS (PT)",
    ("PTBR", 3): "3PTNCNDA - SAMBA INTERM DE NEGOCIOS (PT)",
}

# ── Paleta / tipografia ───────────────────────────────────────────────────────
ORANGE  = "FA8200"
BLACK   = "1A1A1A"
WHITE   = "FFFFFF"
GRAY    = "888888"
LGRAY   = "DDDDDD"
TEXT    = "2D2D2D"
TEXTLT  = "555555"
BGLIGHT = "FFF8F0"
FONT    = "Montserrat"
FONT_SM = Pt(9)
FONT_MD = Pt(10.5)
FONT_LG = Pt(13)

# ── Party I (Samba Export — sempre fixo) ─────────────────────────────────────
PARTY_I = {
    "party_num":   1,
    "party_roman": "I",
    "role":        "Originator",
    "header":      "PARTY I — ORIGINATOR",
    "full_name":   "SAMBA EXPORT LTDA",
    "tax_id":      "60.280.015/0001-82",
    "country":     "Brazil",
    "address":     "Av. Brigadeiro Faria Lima, 1811, Suite 115\nSão Paulo, SP — Brazil",
    "legal_rep":   "Marcelo Soares Magalhães Nogueira",
    "passport":    "N/A (Brazilian Entity)",
    "fixed":       True,
}

_ROLE_LABELS = {
    1: "PARTY I — ORIGINATOR",
    2: "PARTY II — BR INTERMEDIARY",
    3: "PARTY III — BUYER-SIDE",
    4: "PARTY IV — ADDITIONAL",
}


# ═════════════════════════════════════════════════════════════════════════════
# Helpers de formatação (docx) — baseados em generate_ncnda_en.py
# ═════════════════════════════════════════════════════════════════════════════

def _hex_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper().lstrip("#"))
    tcPr.append(shd)


def _cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"),    str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _no_borders(table):
    tbl    = table._tbl
    tblPr  = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _table_width_pct(table, pct: int = 100):
    tbl   = table._tbl
    tblPr = tbl.tblPr
    tblW  = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(pct * 50))
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


def _run(para, text: str, bold: bool = False, italic: bool = False,
         color: str | None = None, size: Pt | None = None) -> object:
    run = para.add_run(text)
    run.font.name  = FONT
    run.font.size  = size or FONT_MD
    run.bold       = bold
    run.italic     = italic
    run.font.color.rgb = _hex_rgb(color or TEXT)
    return run


def _para(doc: Document, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_before=Pt(0), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after  = space_after
    return p


def _blank(doc: Document, n: int = 1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)


def _cell_para(cell, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def _cell_add_line(cell, text: str, bold=False, color=TEXT, size=FONT_MD):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = size; r.bold = bold
    r.font.color.rgb = _hex_rgb(color)


def _clause_label(doc, number: str):
    p = _para(doc, WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(12))
    _run(p, f"CLAUSE {number}", bold=True, color=ORANGE, size=FONT_SM)


def _clause_title(doc, title: str):
    p = _para(doc, WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6))
    _run(p, title, bold=True, size=Pt(13), color=BLACK)


def _body(doc, parts: list, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=0.0):
    p = _para(doc, align)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    for text, bold in parts:
        _run(p, text, bold=bold, size=FONT_MD)
    return p


def _list_line(doc, text: str):
    p = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    _run(p, text, size=FONT_MD)


def _highlight_box(doc, label: str, content: str):
    tbl  = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    _table_width_pct(tbl, 100)
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, BGLIGHT)
    # borda esquerda laranja
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for side in ("top", "right", "bottom", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tcB.append(b)
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:color"), ORANGE)
    tcB.append(left)
    tcPr.append(tcB)
    _cell_margins(cell, top=100, bottom=100, left=150, right=150)

    p1 = _cell_para(cell)
    r1 = p1.add_run(label)
    r1.font.name = FONT; r1.font.size = FONT_MD; r1.bold = True
    r1.font.color.rgb = _hex_rgb(BLACK)
    r2 = p1.add_run(content)
    r2.font.name = FONT; r2.font.size = FONT_MD
    r2.font.color.rgb = _hex_rgb(TEXT)


def _divider(doc):
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tcB.append(b)
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:color"), LGRAY)
    tcB.append(bot)
    tcPr.append(tcB)
    row = tbl.rows[0]
    row.height = Pt(2)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    _table_width_pct(tbl, 100)


# ═════════════════════════════════════════════════════════════════════════════
# Seções do documento
# ═════════════════════════════════════════════════════════════════════════════

def _build_cover(doc: Document, parties: List[dict], today_str: str, ref_code: str):
    tbl  = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    _table_width_pct(tbl, 100)
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, BLACK)
    _cell_margins(cell, top=700, bottom=700, left=700, right=700)
    row = tbl.rows[0]
    row.height = Cm(20)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # Logo
    p_logo = _cell_para(cell)
    r_s = p_logo.add_run("samba")
    r_s.font.name = FONT; r_s.font.size = Pt(22); r_s.bold = True
    r_s.font.color.rgb = _hex_rgb(ORANGE)
    r_e = p_logo.add_run("EXPORT")
    r_e.font.name = FONT; r_e.font.size = Pt(22); r_e.bold = True
    r_e.font.color.rgb = _hex_rgb(WHITE)

    for _ in range(3):
        pb = cell.add_paragraph()
        pb.paragraph_format.space_before = Pt(0)
        pb.paragraph_format.space_after  = Pt(0)

    # Título
    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(2)
    rt = p_title.add_run("Non-Circumvention,\nNon-Disclosure &\n")
    rt.font.name = FONT; rt.font.size = Pt(22); rt.bold = True
    rt.font.color.rgb = _hex_rgb(WHITE)
    rt2 = p_title.add_run("Confidentiality Agreement")
    rt2.font.name = FONT; rt2.font.size = Pt(22); rt2.bold = True
    rt2.font.color.rgb = _hex_rgb(ORANGE)

    p_sub = cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_before = Pt(4)
    p_sub.paragraph_format.space_after  = Pt(20)
    rs = p_sub.add_run("NCNDA — International Commercial Intermediation")
    rs.font.name = FONT; rs.font.size = FONT_SM
    rs.font.color.rgb = _hex_rgb(GRAY)

    p_line = cell.add_paragraph()
    r_line = p_line.add_run("━━━━━━━━━━━━━━━")
    r_line.font.name = FONT; r_line.font.size = Pt(8)
    r_line.font.color.rgb = _hex_rgb(ORANGE)

    # Metadados
    meta = [
        ("DATE OF EXECUTION",      today_str),
        ("DOCUMENT REFERENCE",     ref_code),
        ("PARTY I — ORIGINATOR",   "SAMBA EXPORT LTDA"),
        ("GOVERNING LAW",          "Brazilian Law — São Paulo Courts"),
    ]
    for i, p in enumerate(parties):
        label = _ROLE_LABELS.get(p["party_num"], f"PARTY {p['party_roman']}")
        meta.append((label, p["full_name"]))

    for label, val in meta:
        pb = cell.add_paragraph()
        pb.paragraph_format.space_before = Pt(8)
        pb.paragraph_format.space_after  = Pt(0)
        rl = pb.add_run(label + "\n")
        rl.font.name = FONT; rl.font.size = Pt(7.5); rl.bold = True
        rl.font.color.rgb = _hex_rgb(GRAY)
        rv = pb.add_run(val)
        rv.font.name = FONT; rv.font.size = Pt(11); rv.bold = True
        rv.font.color.rgb = _hex_rgb(WHITE)

    p_conf = cell.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_conf.paragraph_format.space_before = Pt(20)
    rc = p_conf.add_run("CONFIDENTIAL")
    rc.font.name = FONT; rc.font.size = FONT_SM; rc.bold = True
    rc.font.color.rgb = _hex_rgb("333333")

    doc.add_page_break()


def _build_doc_header(doc: Document, ref_code: str, today_str: str):
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    _table_width_pct(tbl, 100)

    for cell in tbl.rows[0].cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcB  = OxmlElement("w:tcBorders")
        for side in ("top", "left", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tcB.append(b)
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "12")
        bot.set(qn("w:color"), ORANGE)
        tcB.append(bot)
        tcPr.append(tcB)

    cell_l = tbl.rows[0].cells[0]
    _cell_margins(cell_l, top=120, bottom=120, left=0, right=0)
    p_l = _cell_para(cell_l)
    r1 = p_l.add_run("samba")
    r1.font.name = FONT; r1.font.size = Pt(16); r1.bold = True
    r1.font.color.rgb = _hex_rgb(ORANGE)
    r2 = p_l.add_run("EXPORT")
    r2.font.name = FONT; r2.font.size = Pt(16); r2.bold = True
    r2.font.color.rgb = _hex_rgb(BLACK)

    cell_r = tbl.rows[0].cells[1]
    _cell_margins(cell_r, top=100, bottom=100, left=0, right=0)
    p_r = _cell_para(cell_r, WD_ALIGN_PARAGRAPH.RIGHT)
    ri = p_r.add_run(
        f"NCNDA — Non-Circumvention, Non-Disclosure & Confidentiality Agreement\n"
        f"Ref: {ref_code}  |  Date: {today_str}"
    )
    ri.font.name = FONT; ri.font.size = Pt(7.5)
    ri.font.color.rgb = _hex_rgb(TEXTLT)

    _blank(doc)


def _build_parties_table(doc: Document, parties: List[dict]):
    """
    Constrói a tabela de partes de forma dinâmica.
    Party I (Samba) sempre na primeira coluna; parties[0..] nas seguintes.
    """
    all_parties = [PARTY_I] + parties  # lista completa (1 + N)
    n_cols = len(all_parties)

    # Campos por papel
    def _fields_for(p: dict) -> list:
        if p.get("fixed"):  # Party I
            return [
                ("LEGAL NAME",              p["full_name"],  True),
                ("BRAZILIAN TAX ID (CNPJ)", p["tax_id"],     True),
                ("REGISTERED ADDRESS",      p["address"],    True),
                ("LEGAL REPRESENTATIVE",    p["legal_rep"],  True),
                ("ROLE IN AGREEMENT",       "Brazilian Commodity Originator & Exporter", True),
            ]
        return [
            ("LEGAL NAME",                   p["full_name"],  False),
            ("COUNTRY OF INCORPORATION",      p["country"],    False),
            ("COMPANY REGISTRATION / TAX ID", p["tax_id"],     False),
            ("REGISTERED ADDRESS",            p["address"],    False),
            ("LEGAL REPRESENTATIVE",          p["legal_rep"],  False),
            ("PASSPORT / ID NUMBER",          p["passport"],   False),
        ]

    field_sets = [_fields_for(p) for p in all_parties]
    max_rows   = max(len(fs) for fs in field_sets)

    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_width_pct(tbl, 100)

    # Bordas finas cinza
    tblB = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:color"), LGRAY)
        tblB.append(b)
    tbl._tbl.tblPr.append(tblB)

    # Cabeçalhos
    hrow = tbl.rows[0]
    for i, p in enumerate(all_parties):
        cell = hrow.cells[i]
        _cell_bg(cell, BLACK)
        _cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p_h = _cell_para(cell)
        r_dot = p_h.add_run("● ")
        r_dot.font.name = FONT; r_dot.font.size = FONT_SM; r_dot.bold = True
        r_dot.font.color.rgb = _hex_rgb(ORANGE)
        label = _ROLE_LABELS.get(p["party_num"], f"PARTY {p['party_roman']}")
        r_lbl = p_h.add_run(label)
        r_lbl.font.name = FONT; r_lbl.font.size = FONT_SM; r_lbl.bold = True
        r_lbl.font.color.rgb = _hex_rgb(WHITE)

    # Linhas de campos
    for fi in range(max_rows):
        row = tbl.add_row()
        for ci, (p, fs) in enumerate(zip(all_parties, field_sets)):
            cell = row.cells[ci]
            _cell_margins(cell, top=100, bottom=80, left=120, right=120)
            if fi < len(fs):
                label, value, is_fixed = fs[fi]
                p_lbl = _cell_para(cell)
                r_l = p_lbl.add_run(label)
                r_l.font.name = FONT; r_l.font.size = Pt(7.5); r_l.bold = True
                r_l.font.color.rgb = _hex_rgb(GRAY)
                p_val = cell.add_paragraph()
                p_val.paragraph_format.space_before = Pt(2)
                p_val.paragraph_format.space_after  = Pt(0)
                r_v = p_val.add_run(value)
                r_v.font.name = FONT; r_v.font.size = FONT_SM; r_v.bold = is_fixed
                r_v.font.color.rgb = _hex_rgb(BLACK if is_fixed else "444444")

    _blank(doc)


def _build_signatures(doc: Document, parties: List[dict]):
    """Bloco de assinaturas dinâmico."""
    all_parties = [PARTY_I] + parties
    n_cols = len(all_parties)

    p_sig_title = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    _run(p_sig_title,
         "SIGNATURES — IN WITNESS WHEREOF, THE PARTIES HAVE EXECUTED THIS AGREEMENT",
         bold=True, size=FONT_SM, color=TEXTLT)

    sig_tbl = doc.add_table(rows=1, cols=n_cols)
    _table_width_pct(sig_tbl, 100)

    for ci, p in enumerate(all_parties):
        cell = sig_tbl.rows[0].cells[ci]
        _cell_margins(cell, top=100, bottom=100, left=80, right=80)

        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcB  = OxmlElement("w:tcBorders")
        for side in ("bottom", "left", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tcB.append(b)
        top_b = OxmlElement("w:top")
        top_b.set(qn("w:val"),   "single")
        top_b.set(qn("w:sz"),    "12")
        top_b.set(qn("w:color"), BLACK)
        tcB.append(top_b)
        tcPr.append(tcB)

        party_label = _ROLE_LABELS.get(p["party_num"], f"PARTY {p['party_roman']}")
        p0 = _cell_para(cell)
        r0 = p0.add_run(party_label)
        r0.font.name = FONT; r0.font.size = Pt(7.5); r0.bold = True
        r0.font.color.rgb = _hex_rgb(ORANGE)

        p1 = cell.add_paragraph()
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after  = Pt(2)
        r1 = p1.add_run(p["full_name"])
        r1.font.name = FONT; r1.font.size = FONT_SM; r1.bold = True
        r1.font.color.rgb = _hex_rgb(BLACK)

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(14)
        rep_text = p.get("legal_rep", "____________________________")
        r2 = p2.add_run(rep_text + "\nLegal Representative")
        r2.font.name = FONT; r2.font.size = FONT_SM
        r2.font.color.rgb = _hex_rgb(TEXTLT)

        p3 = cell.add_paragraph()
        p3.paragraph_format.space_before = Pt(28)
        p3.paragraph_format.space_after  = Pt(2)
        r3 = p3.add_run("_" * 28)
        r3.font.name = FONT; r3.font.size = FONT_SM
        r3.font.color.rgb = _hex_rgb("CCCCCC")

        p4 = cell.add_paragraph()
        p4.paragraph_format.space_before = Pt(2)
        r4 = p4.add_run("Signature & Date")
        r4.font.name = FONT; r4.font.size = Pt(7.5)
        r4.font.color.rgb = _hex_rgb(GRAY)

    _blank(doc, 2)


def _build_footer(doc: Document, ref_code: str):
    ftbl = doc.add_table(rows=1, cols=3)
    _table_width_pct(ftbl, 100)
    _no_borders(ftbl)
    for cell in ftbl.rows[0].cells:
        _cell_bg(cell, BLACK)
        _cell_margins(cell, top=100, bottom=100, left=120, right=120)

    fl = ftbl.rows[0].cells[0]
    pf1 = _cell_para(fl)
    rf1 = pf1.add_run("samba")
    rf1.font.name = FONT; rf1.font.size = Pt(9); rf1.bold = True
    rf1.font.color.rgb = _hex_rgb(ORANGE)
    rf2 = pf1.add_run("EXPORT — CONFIDENTIAL")
    rf2.font.name = FONT; rf2.font.size = Pt(9); rf2.bold = True
    rf2.font.color.rgb = _hex_rgb("888888")

    fc = ftbl.rows[0].cells[1]
    pf2 = _cell_para(fc, WD_ALIGN_PARAGRAPH.CENTER)
    rf3 = pf2.add_run(f"{ref_code}  |  Brazilian Law — São Paulo")
    rf3.font.name = FONT; rf3.font.size = Pt(7.5)
    rf3.font.color.rgb = _hex_rgb("555555")

    fr = ftbl.rows[0].cells[2]
    pf3 = _cell_para(fr, WD_ALIGN_PARAGRAPH.RIGHT)
    rf4 = pf3.add_run("sambaexport.com.br")
    rf4.font.name = FONT; rf4.font.size = Pt(7.5)
    rf4.font.color.rgb = _hex_rgb("555555")


# ═════════════════════════════════════════════════════════════════════════════
# ═════════════════════════════════════════════════════════════════════════════
# Template-based helpers
# ═════════════════════════════════════════════════════════════════════════════

def _build_replacements(parties: List[dict], today_str: str, ref_code: str) -> dict:
    """
    Constrói o dicionário de substituições {{PLACEHOLDER}} → valor.
    Usado quando o documento é gerado a partir de um template Drive.
    """
    r: dict = {
        "{{EXEC_DATE}}": today_str,
        "{{REF_CODE}}":  ref_code,
        "{{TODAY}}":     today_str,
    }
    for i, p in enumerate(parties, start=2):
        r[f"{{{{P{i}_NAME}}}}"]      = p.get("full_name",  "")
        r[f"{{{{P{i}_SHORTNAME}}}}"] = p.get("short_name", "")
        r[f"{{{{P{i}_COUNTRY}}}}"]   = p.get("country",    "")
        r[f"{{{{P{i}_TAXID}}}}"]     = p.get("tax_id",     "")
        r[f"{{{{P{i}_ADDRESS}}}}"]   = p.get("address",    "")
        r[f"{{{{P{i}_REP}}}}"]       = p.get("legal_rep",  "")
        r[f"{{{{P{i}_PASSPORT}}}}"]  = p.get("passport",   "")
    return r


def _replace_in_docx(doc: Document, replacements: dict) -> None:
    """
    Substitui placeholders {{CHAVE}} em todos os runs do documento.
    Estratégia: tenta substituição run a run (preserva formatação).
    Para placeholders que cruzam runs, consolida o texto do parágrafo
    no primeiro run e zera os demais.
    """
    def _replace_para(para) -> None:
        # Verifica se algum placeholder está no texto completo do parágrafo
        full = "".join(r.text for r in para.runs)
        if not any(k in full for k in replacements):
            return

        # Tenta substituição run a run (preserva bold/italic/tamanho)
        for run in para.runs:
            for key, val in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)

        # Verifica se ainda restam placeholders (caso span-runs)
        full2 = "".join(r.text for r in para.runs)
        if any(k in full2 for k in replacements):
            for key, val in replacements.items():
                full2 = full2.replace(key, val)
            # Consolida no primeiro run, zera os demais
            if para.runs:
                para.runs[0].text = full2
                for r in para.runs[1:]:
                    r.text = ""

    for para in doc.paragraphs:
        _replace_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_para(para)


# ═════════════════════════════════════════════════════════════════════════════
# Builder principal (geração programática — fallback quando template ausente)
# ═════════════════════════════════════════════════════════════════════════════

def build_ncnda_document(parties: List[dict], today_str: str, ref_code: str) -> Document:
    """
    Gera o documento NCNDA completo.
    parties: lista de dicts (Party II em diante, SEM Party I)
    """
    doc = Document()

    # Margens
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    _build_cover(doc, parties, today_str, ref_code)
    _build_doc_header(doc, ref_code, today_str)
    _blank(doc)

    # Título
    p = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    _run(p, "Non-Circumvention, Non-Disclosure &\n", bold=True, size=Pt(18), color=BLACK)
    _run(p, "Confidentiality Agreement", bold=True, size=Pt(18), color=ORANGE)
    p2 = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))
    _run(p2, "NCNDA — International Agricultural Commodity Intermediation",
         size=FONT_SM, color=TEXTLT)

    # Preamble
    pp = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(10))
    _run(pp, 'This Non-Circumvention, Non-Disclosure and Confidentiality Agreement (')
    _run(pp, '"Agreement"', bold=True)
    _run(pp, f') is entered into as of {today_str} by and among the following parties (individually a ')
    _run(pp, '"Party"', bold=True)
    _run(pp, ' and collectively the ')
    _run(pp, '"Parties"', bold=True)
    _run(pp, '):')

    _build_parties_table(doc, parties)
    _divider(doc)
    _blank(doc)

    # ── CLAUSE 1 ──────────────────────────────────────────────────────────
    _clause_label(doc, "1")
    _clause_title(doc, "Purpose")
    _body(doc, [
        ("The purpose of this Agreement is to formalize the collaborative intermediation "
         "relationship among the Parties in connection with the purchase, sale, and export of "
         "agricultural commodities of Brazilian origin — including but not limited to soybeans, "
         "soybean meal, soybean oil, corn, sugar (ICUMSA 45, VHP, Crystal), coffee, cotton, and "
         "animal proteins — as well as any other business arising directly or indirectly from "
         "introductions made under this Agreement.", False),
    ])
    _blank(doc)
    _body(doc, [
        ('The Parties acknowledge that each brings to this collaboration confidential information, '
         'business relationships, and strategic contacts (', False),
        ('"Protected Contacts"', True),
        (') that have intrinsic commercial value and must be protected accordingly.', False),
    ])

    # ── CLAUSE 2 ──────────────────────────────────────────────────────────
    _clause_label(doc, "2")
    _clause_title(doc, "Role of the Parties")
    _body(doc, [
        ("Each Party acts exclusively as a ", False),
        ("commercial intermediary", True),
        (" in the transactions facilitated under this Agreement. No Party shall be deemed "
         "a principal buyer, seller, or guarantor of any commodity transaction.", False),
    ])
    _blank(doc)
    _body(doc, [("Accordingly, no Party assumes any liability for:", False)])
    for item in [
        "Product quality, specifications, or conformity;",
        "Logistics, shipping, or customs clearance;",
        "Financial performance or payment by end buyers;",
        "Force majeure events affecting the underlying commodity transaction.",
    ]:
        _list_line(doc, f"• {item}")
    _blank(doc)
    _body(doc, [
        ("All liability for the underlying commercial transaction rests solely with the "
         "parties to the main Sales & Purchase Agreement (SPA).", False),
    ])

    # ── CLAUSE 3 ──────────────────────────────────────────────────────────
    _clause_label(doc, "3")
    _clause_title(doc, "Commission Entitlement")
    _body(doc, [
        ("Each Party shall be entitled to receive a commission for its role in facilitating "
         "transactions under this Agreement. Commission amounts, rates, and split arrangements "
         "shall be agreed on a per-deal basis and may be formalized via:", False),
    ])
    for item in [
        "Written addendum to this Agreement; or",
        "Confirming message exchanged via WhatsApp, Telegram, WeChat, or corporate email — "
        "which shall have the same legal force as a written addendum.",
    ]:
        _list_line(doc, f"• {item}")
    _blank(doc)
    _highlight_box(doc,
        "Payment timing: ",
        "All commissions shall be paid within 1 (one) business day of confirmed receipt of "
        "cleared funds in the designated bank account of the paying Party.")
    _blank(doc)
    _body(doc, [
        ("Each Party is solely responsible for the payment of all applicable taxes, duties, "
         "and levies on its respective commission income under the laws of its jurisdiction.", False),
    ])

    # ── CLAUSE 4 ──────────────────────────────────────────────────────────
    _clause_label(doc, "4")
    _clause_title(doc, "Commission Distribution")
    _body(doc, [
        ("The total gross commission generated from any transaction intermediated under this "
         "Agreement shall be distributed among the Parties according to the percentages or "
         "amounts agreed for each specific deal, as formalized under Clause 3 above.", False),
    ])
    _blank(doc)
    _body(doc, [
        ("No Party shall be entitled to alter the agreed distribution without the written "
         "consent of all Parties.", False),
    ])

    # ── CLAUSE 5 ──────────────────────────────────────────────────────────
    _clause_label(doc, "5")
    _clause_title(doc, "Confidentiality, Non-Circumvention & Non-Compete")

    for num, label, text in [
        ("5.1", "Confidentiality.",
         "Each Party agrees to keep strictly confidential all information received from the "
         "other Parties, including but not limited to: business contacts, pricing, terms of "
         "deals, buyer and seller identities, and internal strategy. Confidential information "
         "shall not be disclosed to any third party without prior written consent."),
        ("5.2", "Non-Circumvention.",
         "Each Party irrevocably undertakes not to directly or indirectly contact, negotiate "
         "with, or transact with any Protected Contact introduced by another Party, without "
         "that Party's prior written authorization and without ensuring that Party's commission "
         "is properly secured."),
    ]:
        p = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
        p.paragraph_format.left_indent = Cm(0.8)
        _run(p, f"{num}  ", bold=True, color=ORANGE)
        _run(p, label + " ", bold=True)
        _run(p, text)
        _blank(doc)

    p53 = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p53.paragraph_format.left_indent = Cm(0.8)
    _run(p53, "5.3  ", bold=True, color=ORANGE)
    _run(p53, "Consequences of Circumvention. ", bold=True)
    _run(p53, "Any Party found to have circumvented another Party shall be liable for:")
    for num2, text2 in [
        ("5.3.1", "Full payment of the commission that would have been due to the circumvented Party;"),
        ("5.3.2", "All direct and indirect damages proven to have resulted from the circumvention;"),
        ("5.3.3", "A contractual penalty to be agreed by the Parties, without prejudice to additional damages."),
    ]:
        p2 = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
        p2.paragraph_format.left_indent = Cm(1.5)
        _run(p2, f"{num2}  ", bold=True, color=ORANGE)
        _run(p2, text2)
    _blank(doc)

    for num, label, text in [
        ("5.4", "Non-Compete.",
         "No Party shall, for a period of 2 (two) years after the termination of this Agreement, "
         "directly or indirectly contact, solicit, or transact with any Protected Contact "
         "introduced under this Agreement, without the introducing Party's prior written consent."),
        ("5.5", "Survival.",
         "The obligations under this Clause 5 shall survive termination of this Agreement for "
         "a period of 5 (five) years (confidentiality and non-circumvention) and "
         "2 (two) years (non-compete)."),
    ]:
        p = _para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
        p.paragraph_format.left_indent = Cm(0.8)
        _run(p, f"{num}  ", bold=True, color=ORANGE)
        _run(p, label + " ", bold=True)
        _run(p, text)
        _blank(doc)

    # ── CLAUSE 6 ──────────────────────────────────────────────────────────
    _clause_label(doc, "6")
    _clause_title(doc, "Governing Law & Jurisdiction")
    _body(doc, [
        ("This Agreement shall be governed by and construed in accordance with the laws of the ", False),
        ("Federative Republic of Brazil", True), (".", False),
    ])
    _blank(doc)
    _body(doc, [
        ("Any dispute arising out of or in connection with this Agreement shall be submitted to "
         "the exclusive jurisdiction of the ", False),
        ("Courts of the City of São Paulo, State of São Paulo, Brazil", True),
        (", with each Party irrevocably waiving any objection to such jurisdiction.", False),
    ])
    _blank(doc)
    _highlight_box(doc, "",
        "The Parties acknowledge that this Agreement may be entered into and enforced across "
        "multiple jurisdictions and that the choice of Brazilian law is deliberate and binding.")
    _blank(doc)

    # ── CLAUSE 7 ──────────────────────────────────────────────────────────
    _clause_label(doc, "7")
    _clause_title(doc, "Term")
    _body(doc, [
        ("This Agreement shall remain in full force and effect for as long as business "
         "operations arising from introductions made hereunder are ongoing. The commission "
         "and confidentiality obligations shall survive any termination or expiry of this "
         "Agreement as set forth in Clause 5.5.", False),
    ])

    # ── CLAUSE 8 ──────────────────────────────────────────────────────────
    _clause_label(doc, "8")
    _clause_title(doc, "Electronic Communications as Binding Agreements")
    _body(doc, [
        ("The Parties expressly agree that messages exchanged via ", False),
        ("WhatsApp, Telegram, WeChat, or corporate email", True),
        (" that confirm commission values, deal terms, or specific amendments to this "
         "Agreement shall constitute legally binding addenda, enforceable as if made in writing.", False),
    ])
    _blank(doc)
    _body(doc, [
        ("No Party shall contest the enforceability of such communications on the grounds "
         "of form or formality.", False),
    ])

    # ── CLAUSE 9 ──────────────────────────────────────────────────────────
    _clause_label(doc, "9")
    _clause_title(doc, "Electronic Signature")
    _body(doc, [
        ("This Agreement may be executed electronically using any of the following platforms: ", False),
        ("DocuSign, ClickSign, Adobe Sign, Gov.br, or ICP-Brasil", True),
        (". Electronically signed versions shall constitute the original and shall have full "
         "legal force and effect under Brazilian law.", False),
    ])
    _blank(doc)
    _body(doc, [("Each Party shall receive a signed copy upon completion of the signature process.", False)])

    _blank(doc)
    _divider(doc)
    _blank(doc)

    _build_signatures(doc, parties)
    _build_footer(doc, ref_code)

    return doc


# ═════════════════════════════════════════════════════════════════════════════
# Agente
# ═════════════════════════════════════════════════════════════════════════════

class NCNDAGeneratorAgent(BaseAgent):
    name = "NCNDAGeneratorAgent"
    description = (
        "Gera NCNDA (Non-Circumvention, Non-Disclosure & Confidentiality Agreement) "
        "de forma programática com número dinâmico de partes. "
        "Upload final como PDF na pasta de saída no Drive."
    )
    visible_in_groups = False
    generates_spreadsheets = False

    def __init__(self, drive: Optional[DriveManager] = None):
        super().__init__()
        self._drive = drive

    @property
    def drive(self) -> DriveManager:
        if self._drive is None:
            self._drive = DriveManager()
        return self._drive

    def process(self, data: Any = None) -> dict:  # noqa: C901
        if not isinstance(data, dict):
            return {"status": "error", "error": "Payload inválido (dict esperado)."}

        parties:         list = data.get("parties", [])
        dry_run:         bool = bool(data.get("dry_run", False))
        language:        str  = (data.get("language") or "EN").upper().strip()
        template_prefix: str  = (data.get("template_prefix") or "").strip()

        if language not in ("EN", "PTBR"):
            language = "EN"
        if not parties:
            return {"status": "error", "error": "Nenhuma parte adicional informada (Party II obrigatória)."}
        if len(parties) > 3:
            return {"status": "error", "error": "Máximo de 3 partes adicionais (Party II, III e IV)."}

        # ── Dados de referência ───────────────────────────────────────────
        today          = datetime.date.today()
        today_str      = today.strftime("%d / %m / %Y")
        today_file     = today.strftime("%Y%m%d")
        party_ii_short = (parties[0].get("short_name") or "PART").upper().replace(" ", "")
        n_additional   = len(parties)
        ref_code       = f"NCNDA-SE-{today_file}-{party_ii_short}"
        filename       = f"{ref_code}.pdf"

        self.log_action("ncnda_start", {
            "ref_code":    ref_code,
            "n_parties":   n_additional + 1,
            "language":    language,
            "dry_run":     dry_run,
        })

        alerts:       list = []
        docx_bytes:   bytes | None = None
        template_used: str = "programmatic"

        # ══════════════════════════════════════════════════════════════════
        # 1. Tentativa: geração via template Drive
        # ══════════════════════════════════════════════════════════════════
        template_name = TEMPLATE_MAP.get((language, n_additional))
        if template_name:
            self.log_action("ncnda_template_search", {"template": template_name})
            try:
                meta = self.drive.find_file_by_name(
                    template_name,
                    TEMPLATES_FOLDER_ID,
                    ignore_underscore_prefix=True,
                )
                if meta:
                    tmpl_bytes = self.drive.fetch_as_docx_bytes(meta)
                    if tmpl_bytes:
                        doc = Document(io.BytesIO(tmpl_bytes))
                        replacements = _build_replacements(parties, today_str, ref_code)
                        _replace_in_docx(doc, replacements)
                        buf = io.BytesIO()
                        doc.save(buf)
                        docx_bytes    = buf.getvalue()
                        template_used = template_name
                        self.log_action("ncnda_template_ok", {"template": template_name})
                    else:
                        alerts.append(f"Template '{template_name}' — falha ao baixar bytes. Usando geração programática.")
                else:
                    alerts.append(f"Template '{template_name}' não encontrado no Drive. Usando geração programática.")
            except Exception as exc:
                alerts.append(f"Erro ao usar template: {exc}. Usando geração programática.")

        # ══════════════════════════════════════════════════════════════════
        # 2. Fallback: geração programática
        # ══════════════════════════════════════════════════════════════════
        if docx_bytes is None:
            try:
                doc = build_ncnda_document(parties, today_str, ref_code)
                buf = io.BytesIO()
                doc.save(buf)
                docx_bytes    = buf.getvalue()
                template_used = "programmatic"
            except Exception as exc:
                return {"status": "error", "error": f"Geração do documento falhou: {exc}"}

        # ── Dry-run ───────────────────────────────────────────────────────
        if dry_run:
            return {
                "status":        "success",
                "dry_run":       True,
                "filename":      filename,
                "size_bytes":    len(docx_bytes),
                "ref_code":      ref_code,
                "language":      language,
                "template_used": template_used,
                "alerts":        alerts,
            }

        # ══════════════════════════════════════════════════════════════════
        # 3. Upload como Google Doc → exporta como PDF
        # ══════════════════════════════════════════════════════════════════
        docx_name = filename.replace(".pdf", ".docx")
        gdoc = self.drive.upload_file_bytes(
            filename=docx_name,
            content=docx_bytes,
            folder_id=OUTPUT_FOLDER_ID,
            save_as_google_doc=True,
        )
        if not gdoc:
            return {"status": "error", "error": "Falha no upload para o Drive.", "filename": filename}

        # Exporta como PDF
        pdf_bytes = self.drive.export_gdoc_as_pdf_bytes(gdoc["id"])
        if pdf_bytes:
            uploaded = self.drive.upload_file_bytes(
                filename=filename,
                content=pdf_bytes,
                folder_id=OUTPUT_FOLDER_ID,
                mime_type="application/pdf",
                save_as_google_doc=False,
            )
            self.drive.delete_file(gdoc["id"])  # remove Google Doc intermediário

            if uploaded:
                self.log_action("ncnda_uploaded", {
                    "file_id":       uploaded.get("id"),
                    "filename":      uploaded.get("name"),
                    "language":      language,
                    "template_used": template_used,
                    "format":        "pdf",
                })
                return {
                    "status":        "success",
                    "file_id":       uploaded.get("id"),
                    "filename":      uploaded.get("name"),
                    "web_link":      uploaded.get("webViewLink"),
                    "size_bytes":    len(pdf_bytes),
                    "file_bytes":    pdf_bytes,
                    "ref_code":      ref_code,
                    "language":      language,
                    "template_used": template_used,
                    "alerts":        alerts,
                }

        # Fallback: mantém Google Doc se conversão PDF falhar
        self.log_action("ncnda_uploaded_gdoc_fallback", {"gdoc_id": gdoc.get("id")})
        alerts.append("PDF não disponível — documento salvo como Google Doc")
        return {
            "status":        "success",
            "file_id":       gdoc.get("id"),
            "filename":      gdoc.get("name"),
            "web_link":      gdoc.get("webViewLink"),
            "size_bytes":    len(docx_bytes),
            "file_bytes":    docx_bytes,
            "ref_code":      ref_code,
            "language":      language,
            "template_used": template_used,
            "alerts":        alerts,
        }
