# -*- coding: utf-8 -*-
"""
agents/cotacao_agent.py
=======================
CotacaoAgent — gera a Cotação Formal Samba Export em PDF.

Pipeline:
  1. Recebe payload com produtos, parâmetros FCL, dados do comprador
  2. Gera documento .docx programaticamente (python-docx)
  3. Faz upload no Drive como Google Doc → exporta PDF
  4. Retorna {filename, web_link, file_bytes, error}

Pasta de saída: mesma do LOI/NCNDA (OUTPUT_FOLDER_ID).
"""
from __future__ import annotations

import io
import datetime
import logging
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from agents.base_agent import BaseAgent
from services.google_drive import DriveManager

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Configurações ──────────────────────────────────────────────────────────────
OUTPUT_FOLDER_ID = "1CPlF_9TtEZ32B5eTAb4b4jRC-h7L10Z2"

# ── Paleta ─────────────────────────────────────────────────────────────────────
ORANGE  = "FA8200"
BLACK   = "1A1A1A"
LGRAY   = "DDDDDD"
MGRAY   = "888888"
BGHEAD  = "FFF3E6"    # fundo cabeçalho de tabela
TEXT    = "2D2D2D"
FONT    = "Montserrat"


# ─── helpers docx ─────────────────────────────────────────────────────────────

def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs) -> None:
    """kwargs: top, bottom, left, right  → dict(sz, color, val)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, props in kwargs.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   props.get("val", "single"))
        el.set(qn("w:sz"),    str(props.get("sz", 4)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), props.get("color", "DDDDDD"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _run(para, text: str, bold=False, size_pt=10, color=TEXT, font=FONT) -> None:
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = _rgb(color)


def _para(doc, text="", bold=False, size_pt=10, color=TEXT,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4) -> Any:
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if text:
        _run(p, text, bold=bold, size_pt=size_pt, color=color)
    return p


def _fmt(n: float) -> str:
    """Formato BR milhar, sem decimais."""
    return f"{n:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _fmt2(n: float) -> str:
    return f"{n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


# ─── Construção do documento ───────────────────────────────────────────────────

def _build_docx(payload: dict) -> bytes:
    category      = payload.get("category", "COTAÇÃO")
    products      = payload.get("products", [])
    fcl_weight    = float(payload.get("fcl_weight", 0))
    fcl_month     = int(payload.get("fcl_month", 0))
    months        = int(payload.get("months", 1))
    tot_mt        = float(payload.get("tot_mt", 0))
    tot_month     = float(payload.get("tot_month", 0))
    tot_total     = float(payload.get("tot_total", 0))
    comm_rate     = float(payload.get("comm_rate", 0.015))
    buyer         = payload.get("buyer", {})
    validity_days = int(payload.get("validity_days", 30))
    incoterm      = payload.get("incoterm", "CIF")
    extra         = payload.get("extra", {})
    today_str     = payload.get("date", datetime.date.today().isoformat())
    ref_code      = _make_ref(buyer.get("name", ""), today_str)

    is_graos = fcl_month == 0   # grãos usa volume total, não FCL

    doc = Document()

    # ── Margens ────────────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    # ── Cabeçalho ─────────────────────────────────────────────────────────────
    _add_header(doc, ref_code, today_str, validity_days)

    # ── Destinatário ──────────────────────────────────────────────────────────
    _add_recipient(doc, buyer, incoterm, category)

    # ── Tabela de produtos ─────────────────────────────────────────────────────
    if is_graos:
        _add_graos_table(doc, products, extra, incoterm)
    else:
        _add_protein_table(doc, products, fcl_weight, fcl_month, months, incoterm)

    # ── Resumo do contrato ─────────────────────────────────────────────────────
    _add_summary(doc, tot_mt, tot_month, tot_total, comm_rate, months,
                 is_graos, fcl_month, len(products))

    # ── Termos e condições ─────────────────────────────────────────────────────
    _add_terms(doc, validity_days, incoterm, extra)

    # ── Assinaturas ────────────────────────────────────────────────────────────
    _add_signatures(doc, buyer)

    # ── Rodapé ─────────────────────────────────────────────────────────────────
    _add_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_ref(buyer_name: str, date_str: str) -> str:
    short = re.sub(r"[^A-Za-z0-9]", "", buyer_name)[:8].upper() or "BUY"
    d = date_str.replace("-", "")
    return f"COT-SE-{d}-{short}"


# ─── Seções do documento ───────────────────────────────────────────────────────

def _add_header(doc: Document, ref_code: str, today_str: str, validity_days: int) -> None:
    # Linha laranja topo
    p_top = doc.add_paragraph()
    pPr = p_top._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:top")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "18")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), ORANGE)
    pBdr.append(bot)
    pPr.append(pBdr)
    p_top.paragraph_format.space_after  = Pt(0)
    p_top.paragraph_format.space_before = Pt(0)

    # Empresa + título em duas colunas via tab
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remover bordas da tabela de cabeçalho
    for row in tbl.rows:
        for cell in row.cells:
            for side in ("top", "bottom", "left", "right", "insideH", "insideV"):
                _set_cell_border(cell, **{side: {"val": "none", "sz": 0, "color": "FFFFFF"}})

    left = tbl.cell(0, 0)
    right = tbl.cell(0, 1)

    # Logo Samba — tenta PNG horizontal; fallback texto laranja
    _LOGO_PATH = Path(__file__).resolve().parent.parent / "logo_samba_horiz.png"
    p_left = left.paragraphs[0]
    p_left.paragraph_format.space_before = Pt(0)
    p_left.paragraph_format.space_after  = Pt(2)
    try:
        r_logo = p_left.add_run()
        r_logo.add_picture(str(_LOGO_PATH), width=Cm(4.4))
    except Exception:
        _run(p_left, "SAMBA EXPORT", bold=True, size_pt=15, color=ORANGE)
    p_left.add_run("\n")
    r2 = p_left.add_run("Samba Intermediação de Negócios Ltda.")
    r2.font.name = FONT; r2.font.size = Pt(8); r2.font.color.rgb = _rgb(MGRAY)
    p_left.add_run("\n")
    r3 = p_left.add_run("CNPJ 60.280.015/0001-82  |  sambaexport.com.br")
    r3.font.name = FONT; r3.font.size = Pt(7.5); r3.font.color.rgb = _rgb(MGRAY)

    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p_right, "COTAÇÃO COMERCIAL", bold=True, size_pt=12, color=BLACK)
    p_right.add_run("\n")
    r4 = p_right.add_run(f"Ref.: {ref_code}")
    r4.font.name = FONT; r4.font.size = Pt(8.5); r4.font.color.rgb = _rgb(MGRAY)
    p_right.add_run("\n")
    r5 = p_right.add_run(f"Data: {_format_date(today_str)}  |  Validade: {validity_days} dias")
    r5.font.name = FONT; r5.font.size = Pt(8); r5.font.color.rgb = _rgb(MGRAY)

    doc.add_paragraph()  # espaço


def _format_date(date_str: str) -> str:
    try:
        d = datetime.date.fromisoformat(date_str)
        meses = ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun",
                 "Jul", "Ago", "Set", "Out", "Nov", "Dez"]
        return f"{d.day:02d} {meses[d.month-1]} {d.year}"
    except Exception:
        return date_str


def _add_recipient(doc: Document, buyer: dict, incoterm: str, category: str) -> None:
    # Seção título
    p_sec = doc.add_paragraph()
    p_sec.paragraph_format.space_before = Pt(2)
    p_sec.paragraph_format.space_after  = Pt(6)
    pPr = p_sec._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), LGRAY)
    pBdr.append(bot)
    pPr.append(pBdr)
    _run(p_sec, "DESTINATÁRIO", bold=True, size_pt=8, color=MGRAY)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for row in tbl.rows:
        for cell in row.cells:
            for side in ("top", "bottom", "left", "right"):
                _set_cell_border(cell, **{side: {"val": "none", "sz": 0, "color": "FFFFFF"}})

    cells = tbl.rows[0].cells

    def _field(cell, label, value):
        p = cell.paragraphs[0]
        r_lbl = p.add_run(f"{label}\n")
        r_lbl.font.name = FONT; r_lbl.font.size = Pt(7.5); r_lbl.font.bold = True
        r_lbl.font.color.rgb = _rgb(MGRAY)
        r_val = p.add_run(value or "—")
        r_val.font.name = FONT; r_val.font.size = Pt(10.5); r_val.font.bold = True
        r_val.font.color.rgb = _rgb(BLACK)

    _field(cells[0], "EMPRESA", buyer.get("name", "—"))
    _field(cells[1], "ATENÇÃO / CONTATO", buyer.get("contact", "—"))
    _field(cells[2], "PAÍS / INCOTERM",
           f"{buyer.get('country', '—')}  ·  {incoterm}")

    doc.add_paragraph()

    # Descrição do produto
    p_cat = doc.add_paragraph()
    p_cat.paragraph_format.space_after = Pt(10)
    _run(p_cat, "Produto:  ", bold=True, size_pt=10, color=BLACK)
    _run(p_cat, category, bold=False, size_pt=10, color=TEXT)


def _add_protein_table(doc: Document, products: list, fcl_weight: float,
                       fcl_month: int, months: int, incoterm: str) -> None:
    _section_title(doc, f"PRODUTOS E PREÇOS ({incoterm})")

    headers = ["PRODUTO", "USD/MT", "MT/MÊS", "VALOR/MÊS (USD)", "TOTAL CONTRATO (USD)"]
    col_w   = [Cm(6.5), Cm(2.4), Cm(2.4), Cm(3.5), Cm(4.0)]

    tbl = doc.add_table(rows=1 + len(products) + 1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabeçalho
    for i, (h, w) in enumerate(zip(headers, col_w)):
        cell = tbl.cell(0, i)
        cell.width = w
        _set_cell_bg(cell, BGHEAD)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.name = FONT; r.font.size = Pt(8); r.font.bold = True
        r.font.color.rgb = _rgb(ORANGE)
        _set_cell_border(cell,
            bottom={"val": "single", "sz": 6, "color": ORANGE})

    # Linhas de produto
    for ri, p in enumerate(products, start=1):
        mt    = fcl_weight * fcl_month
        month = mt * p["price"]
        total = month * months
        vals  = [p["name"], f"$ {_fmt2(p['price'])}", f"{_fmt(mt)} MT",
                 f"$ {_fmt(month)}", f"$ {_fmt(total)}"]
        for ci, (val, w) in enumerate(zip(vals, col_w)):
            cell = tbl.cell(ri, ci)
            cell.width = w
            pg = cell.paragraphs[0]
            pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = pg.add_run(val)
            r.font.name = FONT
            r.font.size = Pt(9.5 if ci == 0 else 9)
            r.font.bold = (ci == 0)
            r.font.color.rgb = _rgb(BLACK if ci == 0 else TEXT)
            _set_cell_border(cell,
                bottom={"val": "single", "sz": 2, "color": LGRAY})

    # Total geral
    tot_mt    = sum(fcl_weight * fcl_month for _ in products)
    tot_month = sum(fcl_weight * fcl_month * p["price"] for p in products)
    tot_total = tot_month * months
    ri_tot    = len(products) + 1
    tot_vals  = ["TOTAL GERAL", "", f"{_fmt(tot_mt)} MT",
                 f"$ {_fmt(tot_month)}", f"$ {_fmt(tot_total)}"]
    for ci, (val, w) in enumerate(zip(tot_vals, col_w)):
        cell = tbl.cell(ri_tot, ci)
        cell.width = w
        _set_cell_bg(cell, "F9F9F9")
        pg = cell.paragraphs[0]
        pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if val:
            r = pg.add_run(val)
            r.font.name = FONT; r.font.size = Pt(9.5); r.font.bold = True
            r.font.color.rgb = _rgb(ORANGE if ci in (3, 4) else BLACK)
        _set_cell_border(cell,
            top={"val": "single", "sz": 6, "color": MGRAY})

    doc.add_paragraph()


def _add_graos_table(doc: Document, products: list, extra: dict, incoterm: str) -> None:
    _section_title(doc, f"PREÇO ESTIMADO ({incoterm})")

    commodity = extra.get("commodity", "Grão")
    destino   = extra.get("destino", "China")
    rota      = extra.get("rota", "")
    cambio    = extra.get("cambio", 0.0)

    if not products:
        return

    p = products[0]
    fob = p.get("fob", 0.0)
    cif = p.get("cif", p.get("price", 0.0))
    vol = p.get("volume_mt", 0.0)
    cbot = p.get("cbot", 0.0)
    basis = p.get("basis", 0.0)
    frete = p.get("frete_maritimo", 0.0)
    total = cif * vol

    rows = [
        ("Commodity",        commodity),
        ("Destino",          destino),
        ("Rota de Saída",    rota or "—"),
        ("CBOT Futuro",      f"{_fmt2(cbot)} ¢/bu"),
        ("Basis Porto",      f"{'+' if basis >= 0 else ''}{_fmt2(basis)} ¢/bu"),
        ("FOB Estimado",     f"USD {_fmt2(fob)}/MT"),
        ("Frete Marítimo",   f"USD {_fmt2(frete)}/MT"),
        (f"CIF {destino}",   f"USD {_fmt2(cif)}/MT"),
        ("Volume",           f"{_fmt(vol)} MT"),
        ("Câmbio BRL/USD",   f"R$ {_fmt2(cambio)}"),
        ("VALOR TOTAL",      f"USD {_fmt(total)}"),
    ]

    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    col_w = [Cm(5.5), Cm(8.0)]

    for ri, (label, value) in enumerate(rows):
        is_total = label == "VALOR TOTAL"
        for ci, text in enumerate([label, value]):
            cell = tbl.cell(ri, ci)
            cell.width = col_w[ci]
            if is_total:
                _set_cell_bg(cell, BGHEAD)
            pg = cell.paragraphs[0]
            pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci == 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = pg.add_run(text)
            r.font.name = FONT
            r.font.size = Pt(9.5)
            r.font.bold = is_total or ci == 0
            r.font.color.rgb = _rgb(ORANGE if is_total else (MGRAY if ci == 0 else BLACK))
            _set_cell_border(cell,
                bottom={"val": "single", "sz": 2,
                        "color": ORANGE if is_total else LGRAY})

    doc.add_paragraph()


def _section_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), ORANGE)
    pBdr.append(bot)
    pPr.append(pBdr)
    _run(p, title, bold=True, size_pt=8.5, color=ORANGE)


def _add_summary(doc: Document, tot_mt: float, tot_month: float, tot_total: float,
                 comm_rate: float, months: int, is_graos: bool,
                 fcl_month: int, n_products: int) -> None:
    _section_title(doc, "RESUMO DO CONTRATO")

    comm_total = tot_total * comm_rate
    comm_month = (tot_total / months * comm_rate) if months > 0 else 0.0

    if is_graos:
        summary_items = [
            ("Volume Total",        f"{_fmt(tot_mt)} MT"),
            ("Valor Total (CIF)",   f"USD {_fmt(tot_total)}"),
            (f"Comissão ({comm_rate*100:.1f}%)",
                                    f"USD {_fmt(comm_total)}"),
        ]
    else:
        total_fcl = n_products * fcl_month * months
        summary_items = [
            ("Duração do Contrato", f"{months} meses"),
            ("Volume Mensal",       f"{_fmt(tot_mt)} MT  ({n_products * fcl_month} FCL/mês)"),
            ("Volume Total",        f"{_fmt(tot_mt * months)} MT  ({total_fcl} FCL)"),
            ("Receita Mensal",      f"USD {_fmt(tot_month)}"),
            ("Total do Contrato",   f"USD {_fmt(tot_total)}"),
            (f"Comissão ({comm_rate*100:.1f}%)",
                                    f"USD {_fmt(comm_total)}"),
            ("Comissão/Mês",        f"USD {_fmt(comm_month)}"),
        ]

    tbl = doc.add_table(rows=len(summary_items), cols=2)
    tbl.style = "Table Grid"
    col_w = [Cm(6.5), Cm(6.0)]

    for ri, (label, value) in enumerate(summary_items):
        is_comm = "Comissão" in label and "Mês" not in label
        for ci, text in enumerate([label, value]):
            cell = tbl.cell(ri, ci)
            cell.width = col_w[ci]
            if is_comm:
                _set_cell_bg(cell, BGHEAD)
            pg = cell.paragraphs[0]
            pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci == 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = pg.add_run(text)
            r.font.name = FONT
            r.font.size = Pt(9.5)
            r.font.bold = is_comm
            r.font.color.rgb = _rgb(ORANGE if is_comm else (MGRAY if ci == 0 else BLACK))
            _set_cell_border(cell,
                bottom={"val": "single", "sz": 2,
                        "color": ORANGE if is_comm else LGRAY})

    doc.add_paragraph()


def _add_terms(doc: Document, validity_days: int, incoterm: str, extra: dict) -> None:
    _section_title(doc, "CONDIÇÕES GERAIS")

    rota = extra.get("rota", "")
    rota_note = (
        f" Rota de saída: {rota}." if rota else ""
    )

    terms = [
        f"Validade desta cotação: {validity_days} dias corridos a partir da data de emissão.",
        f"Preços em USD, base {incoterm}, conforme especificação acima.{rota_note}",
        "Pagamento: a combinar — carta de crédito (L/C) ou transferência antecipada (T/T).",
        "Quantidade mínima por pedido: conforme tabela acima. Volumes customizados sob consulta.",
        "Frigorífico/exportador habilitado pelo MAPA e aprovado pelo GACC (China).",
        "Esta cotação é uma estimativa de referência. Valores finais sujeitos a confirmação "
        "em contrato formal após negociação.",
        "Câmbio sujeito a variação. Preços de grãos atrelados a CBOT na data de fixação.",
        "Dúvidas: comercial@sambaexport.com.br",
    ]

    for i, t in enumerate(terms, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, f"{i}.  {t}", size_pt=8.5, color=TEXT)

    doc.add_paragraph()


def _add_signatures(doc: Document, buyer: dict) -> None:
    _section_title(doc, "ASSINATURAS")

    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    col_w = [Cm(8.5), Cm(8.5)]
    sides = ["Samba Export — Emitente", buyer.get("name", "Comprador — Destinatário")]

    for ci, label in enumerate(sides):
        # nome
        cell_lbl = tbl.cell(0, ci)
        cell_lbl.width = col_w[ci]
        for side in ("top", "bottom", "left", "right"):
            _set_cell_border(cell_lbl, **{side: {"val": "none", "sz": 0, "color": "FFFFFF"}})
        pg = cell_lbl.paragraphs[0]
        pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(pg, label, bold=True, size_pt=9, color=BLACK)

        # linha de assinatura
        cell_sig = tbl.cell(1, ci)
        cell_sig.width = col_w[ci]
        for side in ("top", "left", "right"):
            _set_cell_border(cell_sig, **{side: {"val": "none", "sz": 0, "color": "FFFFFF"}})
        _set_cell_border(cell_sig,
            bottom={"val": "single", "sz": 4, "color": MGRAY})
        pg2 = cell_sig.paragraphs[0]
        pg2.paragraph_format.space_before = Pt(22)
        pg2.paragraph_format.space_after  = Pt(2)

        # legenda
        cell_leg = tbl.cell(2, ci)
        cell_leg.width = col_w[ci]
        for side in ("top", "bottom", "left", "right"):
            _set_cell_border(cell_leg, **{side: {"val": "none", "sz": 0, "color": "FFFFFF"}})
        pg3 = cell_leg.paragraphs[0]
        pg3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(pg3, "Nome / Cargo / Data", size_pt=7.5, color=MGRAY)

    doc.add_paragraph()


def _add_footer(doc: Document) -> None:
    p_bot = doc.add_paragraph()
    p_bot.paragraph_format.space_before = Pt(12)
    pPr = p_bot._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "4")
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), LGRAY)
    pBdr.append(top)
    pPr.append(pBdr)
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_bot,
         "Samba Intermediação de Negócios Ltda.  |  CNPJ 60.280.015/0001-82  |  "
         "sambaexport.com.br  |  comercial@sambaexport.com.br",
         size_pt=7.5, color=MGRAY)
    p_dis = doc.add_paragraph()
    p_dis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dis.paragraph_format.space_before = Pt(2)
    _run(p_dis,
         "⚠ Documento de caráter informativo. Não constitui oferta vinculante. "
         "Valores sujeitos a confirmação contratual.",
         size_pt=7, color=MGRAY)


# ─── Upload Drive ─────────────────────────────────────────────────────────────

def _upload_and_export(docx_bytes: bytes, filename_base: str) -> dict:
    """
    Faz upload do .docx como Google Doc → exporta como PDF → deleta intermediário.
    Retorna {file_id, web_link, file_bytes}.
    """
    drive = DriveManager()
    DOCX_MIME = ("application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.document")

    # 1. Upload .docx → converte para Google Doc na nuvem
    gdoc_meta = drive.upload_file_bytes(
        filename=f"{filename_base}.docx",
        content=docx_bytes,
        folder_id=OUTPUT_FOLDER_ID,
        mime_type=DOCX_MIME,
        save_as_google_doc=True,
    )
    gdoc_id = (gdoc_meta or {}).get("id", "")

    # 2. Exporta como PDF
    pdf_bytes = drive.export_gdoc_as_pdf_bytes(gdoc_id) if gdoc_id else None
    if not pdf_bytes:
        # fallback: salva o próprio docx
        return {"file_id": gdoc_id, "web_link": "", "file_bytes": docx_bytes}

    # 3. Deleta o Google Doc intermediário
    try:
        drive.delete_file(gdoc_id)
    except Exception:
        pass

    # 4. Re-upload do PDF final
    pdf_meta = drive.upload_file_bytes(
        filename=f"{filename_base}.pdf",
        content=pdf_bytes,
        folder_id=OUTPUT_FOLDER_ID,
        mime_type="application/pdf",
        save_as_google_doc=False,
    )
    pdf_id   = (pdf_meta or {}).get("id", "")
    web_link = f"https://drive.google.com/file/d/{pdf_id}/view" if pdf_id else ""
    return {"file_id": pdf_id, "web_link": web_link, "file_bytes": pdf_bytes}


# ─── Entry point público ───────────────────────────────────────────────────────

def process_cotacao(payload: dict) -> dict:
    """
    Ponto de entrada chamado pelo cotacao_widget.py.

    payload keys:
      category, products, fcl_weight, fcl_month, months,
      tot_mt, tot_month, tot_total, comm_rate,
      buyer {name, contact, country}, validity_days, incoterm,
      extra {commodity, destino, rota, ...}, date

    Retorna:
      {filename, web_link, file_bytes, error}
    """
    try:
        buyer_name = payload.get("buyer", {}).get("name", "BUYER")
        today_str  = payload.get("date", datetime.date.today().isoformat())
        ref_code   = _make_ref(buyer_name, today_str)
        filename   = f"{ref_code}.pdf"

        # Gera docx
        docx_bytes = _build_docx(payload)

        # Upload Drive
        result = _upload_and_export(docx_bytes, ref_code)

        return {
            "filename":   filename,
            "web_link":   result.get("web_link", ""),
            "file_bytes": result.get("file_bytes"),
            "error":      None,
        }

    except Exception as exc:
        return {
            "filename":   "",
            "web_link":   "",
            "file_bytes": None,
            "error":      str(exc),
        }


# ══════════════════════════════════════════════════════════════════════════════
# PRICE INDICATION — injeção no template real .docx
# Template: 2_PRICE_PREQUOTATION_SOY_SUGAR_CORN.docx
# ══════════════════════════════════════════════════════════════════════════════

_PI_TEMPLATE_LOCAL = Path(
    r"H:\Drives compartilhados\SAMBA EXPORT\MODELOS DE DOCUMENTOS"
    r"\_PRE LOI\2_PRICE_PREQUOTATION_SOY_SUGAR_CORN.docx"
)
_PI_TEMPLATES_FOLDER_ID  = "1EU0KkSzHKhxqOlp3XGC-xvs6Xzvf_XTh"   # pasta _PRE LOI
_PI_TEMPLATE_FILE_ID     = "1Om7bIHUPqcBF817kt3ypF8UJiYRgV2Kv"    # ID direto do .docx


def process_price_indication(payload: dict) -> dict:
    """
    Injeta os marcadores do payload no template real
    2_PRICE_PREQUOTATION_SOY_SUGAR_CORN.docx e gera o PDF final.

    payload esperado (construído por _render_formal_quote_form):
        document_type   : "PRICE_INDICATION"
        template_name   : "2_PRICE_PREQUOTATION_SOY_SUGAR_CORN.docx"
        dynamic_fields  : {
            COMODITIE_TYPE, CITY, Port, COUNTRY,
            DD/MM/YYYY, MM, YYYY,
            FIRST NAME Company, FULL NAME Company, PORTO
        }
        financial_fields: {
            PRICE BASIS - USD XXX,XX,   BASIS REFERENCIA PORTO,
            PRICE FOB USD,              PRICE FREIGHT USD,
            FINAL_PRICE USD,            TOTAL_COMISSION_CONTRACT_USD
        }
        dry_run         : bool  (default False)
        output_format   : "pdf" | "docx"  (default "pdf")

    Retorna: { filename, web_link, file_bytes, error }
    """
    try:
        dyn   = payload.get("dynamic_fields")  or {}
        fin   = payload.get("financial_fields") or {}
        dry   = bool(payload.get("dry_run", False))
        fmt   = (payload.get("output_format") or "pdf").lower()
        tname = payload.get("template_name",
                            "2_PRICE_PREQUOTATION_SOY_SUGAR_CORN.docx")

        company_full = dyn.get("FULL NAME Company", "").strip()
        if not company_full:
            return {"filename": "", "web_link": "", "file_bytes": None,
                    "error": "FULL NAME Company é obrigatório."}

        # 1. Carregar template
        tpl_bytes = _pi_load_template(tname)
        if not tpl_bytes:
            errs = "; ".join(_PI_LAST_LOAD_ERRORS) if _PI_LAST_LOAD_ERRORS else "sem detalhes"
            return {"filename": "", "web_link": "", "file_bytes": None,
                    "error": f"Template '{tname}' nao encontrado. Tentativas: {errs}"}

        # 2. Montar tabela de substituições
        subs = _pi_build_subs(dyn, fin)

        # 3. Renderizar (merge runs + replace + typo fix)
        try:
            docx_bytes = _pi_render(tpl_bytes, subs)
        except Exception as exc:
            return {"filename": "", "web_link": "", "file_bytes": None,
                    "error": f"Renderização falhou: {exc}"}

        # 4. Nome do arquivo base
        first   = dyn.get("FIRST NAME Company",
                           company_full.split()[0])[:15].upper()
        comm    = re.sub(r"[^A-Z0-9]", "_",
                         dyn.get("COMODITIE_TYPE", "COMMODITY").upper())
        today_s = datetime.date.today().strftime("%Y%m%d")
        ref_code = f"PI_{first}_{comm}_{today_s}"

        if dry:
            return {"filename": f"{ref_code}.docx", "web_link": "",
                    "file_bytes": docx_bytes, "error": None, "pdf_failed": False}

        # 5. Upload via Drive (GDoc → PDF) — mesma rota das proteínas, funciona no Cloud
        try:
            up = _upload_and_export(docx_bytes, ref_code)
            return {"filename": f"{ref_code}.pdf", "web_link": up.get("web_link", ""),
                    "file_bytes": up.get("file_bytes"), "error": None, "pdf_failed": False}
        except Exception as exc:
            logger.warning("Upload PI falhou: %s", exc)
            return {"filename": f"{ref_code}.docx", "web_link": "",
                    "file_bytes": docx_bytes, "error": None, "pdf_failed": True}

    except Exception as exc:
        logger.exception("process_price_indication inesperado")
        return {"filename": "", "web_link": "", "file_bytes": None,
                "error": str(exc)}


# ── Helpers privados do PI ─────────────────────────────────────────────────────

def _pi_load_template(template_name: str) -> Optional[bytes]:
    """
    Carrega template PI em 3 tentativas:
      1. Busca por nome na pasta Drive correta
      2. Acesso direto pelo file ID (bypassa busca — robusto no Streamlit Cloud)
      3. Fallback local (ambiente Windows com drive mapeado)

    Em caso de falha total, deixa rastro detalhado no logger e em
    _PI_LAST_LOAD_ERRORS para o widget exibir ao usuário.
    """
    global _PI_LAST_LOAD_ERRORS
    _PI_LAST_LOAD_ERRORS = []

    try:
        from services.google_drive import DriveManager
        drive = DriveManager()
        if not getattr(drive, "service", None):
            _PI_LAST_LOAD_ERRORS.append("Drive service nao inicializado (token expirado no Streamlit Cloud?)")
            logger.warning("Drive service nao inicializado para template PI")

        # 1. Busca por nome na pasta
        if template_name and getattr(drive, "service", None):
            try:
                meta = drive.find_file_by_name(
                    template_name, _PI_TEMPLATES_FOLDER_ID,
                    ignore_underscore_prefix=True,
                )
                if meta:
                    b = drive.fetch_as_docx_bytes(meta)
                    if b:
                        logger.info("Template PI carregado do Drive (busca): %s", template_name)
                        return b
                    else:
                        _PI_LAST_LOAD_ERRORS.append(f"find_file_by_name OK mas fetch_as_docx_bytes retornou vazio")
                else:
                    _PI_LAST_LOAD_ERRORS.append(f"find_file_by_name nao achou '{template_name}' em {_PI_TEMPLATES_FOLDER_ID}")
            except Exception as exc:
                _PI_LAST_LOAD_ERRORS.append(f"find_file_by_name erro: {exc}")
                logger.warning("PI find_file_by_name falhou: %s", exc)

        # 2. ID direto (fallback quando busca falha no Cloud)
        if getattr(drive, "service", None):
            try:
                file_meta = drive.service.files().get(
                    fileId=_PI_TEMPLATE_FILE_ID,
                    fields="id,name,mimeType",
                    supportsAllDrives=True,
                ).execute()
                b = drive.fetch_as_docx_bytes(file_meta)
                if b:
                    logger.info("Template PI carregado do Drive (ID direto): %s", _PI_TEMPLATE_FILE_ID)
                    return b
                else:
                    _PI_LAST_LOAD_ERRORS.append(f"files().get(ID direto) OK mas fetch vazio")
            except Exception as exc:
                _PI_LAST_LOAD_ERRORS.append(f"files().get(ID direto) erro: {exc}")
                logger.warning("Template PI por ID falhou: %s", exc)

    except Exception as exc:
        _PI_LAST_LOAD_ERRORS.append(f"DriveManager() erro: {exc}")
        logger.warning("Drive indisponivel para template PI: %s", exc)

    # 3. Fallback local
    try:
        if _PI_TEMPLATE_LOCAL.exists():
            logger.info("Template PI local: %s", _PI_TEMPLATE_LOCAL)
            return _PI_TEMPLATE_LOCAL.read_bytes()
        else:
            _PI_LAST_LOAD_ERRORS.append(f"Fallback local nao existe: {_PI_TEMPLATE_LOCAL}")
    except Exception as exc:
        _PI_LAST_LOAD_ERRORS.append(f"Fallback local erro: {exc}")

    return None


# Buffer de erros de _pi_load_template — consultado pelo widget para diagnostico
_PI_LAST_LOAD_ERRORS: list = []


def _pi_build_subs(dyn: dict, fin: dict) -> dict:
    """
    Monta a tabela de substituições {MARCADOR} → valor.
    Template V3: (Port} removido — apenas {CITY} e {COUNTRY} para o incoterm.
    """
    return {
        # — dinâmicos —
        "COMODITIE_TYPE":             dyn.get("COMODITIE_TYPE", ""),
        "CITY":                       dyn.get("CITY", "Main Port"),
        "COUNTRY":                    dyn.get("COUNTRY", ""),
        "DD/MM/YYYY":                 dyn.get("DD/MM/YYYY",
                                              datetime.date.today().strftime("%d/%m/%Y")),
        "MM":                         dyn.get("MM",
                                              datetime.date.today().strftime("%m")),
        "YYYY":                       dyn.get("YYYY",
                                              str(datetime.date.today().year)),
        "FIRST NAME Company":         dyn.get("FIRST NAME Company", ""),
        "FULL NAME Company":          dyn.get("FULL NAME Company", ""),
        "PORTO":                      dyn.get("PORTO", ""),
        # — financeiros (nomes EXATOS do XML — atenção a espaços e ausência de "USD") —
        "PRICE BASIS ":               fin.get("PRICE BASIS", ""),       # ← trailing space no XML
        "BASIS REFERENCIA PORTO":     fin.get("BASIS REFERENCIA PORTO", ""),
        "PRICE  FOB":                 fin.get("PRICE FOB", ""),         # ← duplo espaço, sem " USD"
        "PRICE FREIGHT":              fin.get("PRICE FREIGHT", ""),     # ← sem " USD"
        "FINAL_PRICE":                fin.get("FINAL_PRICE", ""),       # ← marcador simples (V2 template)
        "COMISSION_CONTRACT":         fin.get("COMISSION_CONTRACT", "5,00"),
        "QUANTITY_MT ":               fin.get("QUANTITY_MT", ""),       # ← trailing space no XML
    }


def _pi_render(tpl_bytes: bytes, subs: dict) -> bytes:
    """
    Renderiza o template injetando {MARCADORES} → valores.

    ESTRATÉGIA RAW-ZIP (preserva layout de 1 página):
      • Abre o .docx como ZIP sem passar pelo python-docx Document.save().
      • python-docx.save() reescreve o ZIP inteiro e altera propriedades
        de posicionamento das text boxes → quebra em 2 páginas.
      • Aqui: apenas word/document.xml é modificado; todos os outros arquivos
        (settings.xml, styles.xml, numbering.xml, fontes, imagens…) são
        copiados byte a byte do template original.

    LÓGICA DE SUBSTITUIÇÃO por parágrafo:
      • Itera body.iter(_TAG_P) → encontra <w:p> em qualquer profundidade
        (corpo, tabelas, text boxes, txbxContent, mc:Choice, mc:Fallback).
      • Coleta <w:t> APENAS de <w:r> filhos DIRETOS do <w:p> — evita
        descer para text boxes aninhados dentro do âncora do parágrafo.
      • Concatena os textos, substitui marcadores, coloca resultado no
        primeiro <w:t> e zera os demais.
    """
    import zipfile
    from lxml import etree as _ET

    _W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _XML = "http://www.w3.org/XML/1998/namespace"
    _TAG_P = f"{{{_W}}}p"
    _TAG_R = f"{{{_W}}}r"
    _TAG_T = f"{{{_W}}}t"

    def _sub(text: str) -> str:
        for k, v in subs.items():
            text = text.replace(f"{{{k}}}", str(v))
        return text

    def _process_para(p_elem) -> None:
        """Substitui marcadores nos <w:t> de <w:r> filhos diretos do <w:p>."""
        t_nodes = []
        for child in p_elem:           # filhos DIRETOS do <w:p>
            if child.tag == _TAG_R:
                t = child.find(_TAG_T)
                if t is not None:
                    t_nodes.append(t)
        if not t_nodes:
            return

        full = "".join(t.text or "" for t in t_nodes)

        has = any(f"{{{k}}}" in full for k in subs)
        if not has:
            return

        new_text = _sub(full)
        if new_text == full:
            return

        t_nodes[0].text = new_text
        t_nodes[0].set(f"{{{_XML}}}space", "preserve")
        for t in t_nodes[1:]:
            t.text = ""

    def _process_xml(xml_bytes: bytes) -> bytes:
        """Parseia XML com lxml, aplica substituições, serializa de volta."""
        root = _ET.fromstring(xml_bytes)
        for p_elem in root.iter(_TAG_P):
            _process_para(p_elem)
        return _ET.tostring(
            root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

    # Arquivos XML com marcadores (footer não tem marcadores neste template)
    _TARGETS = {"word/document.xml"}

    in_buf  = io.BytesIO(tpl_bytes)
    out_buf = io.BytesIO()

    with zipfile.ZipFile(in_buf, "r") as zin, \
         zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in _TARGETS:
                data = _process_xml(data)
            zout.writestr(item, data)

    return out_buf.getvalue()


def _pi_to_pdf(docx_bytes: bytes) -> Optional[bytes]:
    """
    Converte .docx → .pdf.
    Ordem:
      1) docx2pdf via subprocess isolado (evita bloqueios COM no Streamlit)
      2) docx2pdf in-process (fallback)
      3) LibreOffice headless
      4) None

    Usa pasta fixa em vez de TemporaryDirectory para evitar PermissionError
    no Windows quando Word mantém o arquivo aberto durante a conversão.
    """
    import sys as _sys
    # Pasta de trabalho fixa (evita PermissionError do Windows no cleanup)
    tmp_dir = Path(tempfile.gettempdir()) / "samba_pi_conv"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    src = str(tmp_dir / "pi.docx")
    dst = str(tmp_dir / "pi.pdf")

    try:
        with open(src, "wb") as f:
            f.write(docx_bytes)

        # 1. docx2pdf via subprocess isolado — mais robusto no contexto de servidor
        #    (evita problemas de COM threading quando Streamlit roda como processo de serviço)
        try:
            if os.path.exists(dst):
                os.remove(dst)
            res = subprocess.run(
                [_sys.executable, "-c",
                 "import sys; from docx2pdf import convert; convert(sys.argv[1], sys.argv[2])",
                 src, dst],
                capture_output=True, timeout=120,
            )
            if os.path.exists(dst):
                pdf = open(dst, "rb").read()
                logger.info("PI PDF via docx2pdf subprocess: %d bytes", len(pdf))
                return pdf
            logger.warning(
                "docx2pdf subprocess: rc=%d stderr=%s",
                res.returncode,
                res.stderr.decode(errors="replace")[:300],
            )
        except Exception as exc:
            logger.warning("docx2pdf subprocess falhou: %s", exc)

        # 2. docx2pdf in-process (fallback caso subprocess não funcione)
        try:
            from docx2pdf import convert as _d2p
            if os.path.exists(dst):
                os.remove(dst)
            _d2p(src, dst)
            if os.path.exists(dst):
                pdf = open(dst, "rb").read()
                logger.info("PI PDF via docx2pdf in-process: %d bytes", len(pdf))
                return pdf
        except Exception as exc:
            logger.warning("docx2pdf in-process falhou: %s", exc)

        # 3. LibreOffice headless
        soffice = _pi_find_soffice()
        if soffice:
            try:
                res = subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf",
                     "--outdir", str(tmp_dir), src],
                    capture_output=True, timeout=60,
                )
                if res.returncode == 0 and os.path.exists(dst):
                    pdf = open(dst, "rb").read()
                    logger.info("PI PDF via LibreOffice: %d bytes", len(pdf))
                    return pdf
                logger.warning("LibreOffice rc=%d", res.returncode)
            except Exception as exc:
                logger.warning("LibreOffice falhou: %s", exc)

    except Exception as exc:
        logger.warning("_pi_to_pdf erro geral: %s", exc)
    finally:
        # Limpeza gentil — ignora se Word ainda tem o arquivo aberto
        for f in (src, dst):
            try:
                if os.path.exists(f):
                    os.remove(f)
            except Exception:
                pass

    logger.warning("Nenhum conversor PDF disponível — PI retornará .docx")
    return None


def _pi_find_soffice() -> Optional[str]:
    for c in [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]:
        if os.path.isfile(c):
            return c
    try:
        subprocess.run(["soffice", "--version"], capture_output=True, timeout=5)
        return "soffice"
    except Exception:
        return None


# ─── Classe agent (para compatibilidade com base_agent) ───────────────────────

class CotacaoAgent(BaseAgent):
    name        = "CotacaoAgent"
    description = (
        "Gera Cotação Formal Samba Export em PDF para proteínas (frango/suína) "
        "e grãos (soja — incluindo Arco Norte, milho, açúcar). "
        "Upload automático no Google Drive."
    )
    visible_in_groups      = False
    generates_spreadsheets = False

    def process(self, payload: dict | None = None, **kwargs) -> dict:
        if payload is None:
            payload = kwargs
        return process_cotacao(payload)
